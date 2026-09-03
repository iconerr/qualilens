# Copyright 2026 Ashita Aggarwal and Suraj Commuri
# SPDX-License-Identifier: Apache-2.0

"""Regression tests for the defects found in the audit: path traversal,
resume duplication, cancellation, checkpoint races, merge-chain handling,
researcher-edit preservation, deletion cleanup, restart reconciliation,
truncation refusal, and quote normalization."""

import io
import json
import threading
import time

import pytest
from starlette.testclient import TestClient

import app.db as db
import app.llm as llm
import app.pipeline as pipeline
from app.main import app, SESSION_TOKEN
from app.methods.base import RunContext, apply_code_review_resolution, locate_quote
from app.methods import common

# every request must look like it comes from the app itself: a local Host
# and the per-launch session token (see main.local_only_guard)
AUTH = {"X-QualiLens-Token": SESSION_TOKEN}
client = TestClient(app, base_url="http://127.0.0.1", headers=AUTH)

DOC_A = "We chose the vendor because the price was transparent. Support was responsive."
DOC_B = "The onboarding was slow at first. The transparent pricing convinced finance."
DOC_C = "Contract terms were flexible. Support answered within an hour every time."


# ---------- mock LLM machinery ----------

class MockLLM:
    """Scriptable llm.chat replacement: counts calls, can fail selected calls
    once, and returns canned JSON keyed on prompt content."""

    def __init__(self):
        self.calls = []
        self.fail_once_when = None   # tuple of substrings that ALL must match, fails once
        self._failed = False
        self.on_call = None          # hook(user) called before answering

    def __call__(self, provider, model, api_key, system, user,
                 max_tokens=8000, temperature=0.3):
        self.calls.append(user[:80])
        if self.on_call:
            self.on_call(user)
        if (self.fail_once_when and not self._failed
                and all(s in user for s in self.fail_once_when)):
            self._failed = True
            raise llm.LLMError("simulated transient failure")
        usage = {"input_tokens": 100, "output_tokens": 50, "stop_reason": "end_turn"}
        return self._answer(system, user), usage

    def _answer(self, system, user):
        if "deriving a CODEBOOK" in system:
            return json.dumps({"codes": [
                {"name": "Pricing", "definition": "Price talk.", "inclusion_criteria": "",
                 "example": "the price was transparent"},
                {"name": "Support", "definition": "Support talk.", "inclusion_criteria": "",
                 "example": "Support was responsive"}]})
        if "APPLYING a fixed codebook" in system:
            if "price was transparent" in user:
                return json.dumps({"assignments": [
                    {"code": "Pricing", "quote": "the price was transparent", "confidence": 0.9},
                    {"code": "support", "quote": "Support was responsive", "confidence": 0.9},
                    {"code": "Nonexistent Code", "quote": "Support was responsive", "confidence": 0.9}]})
            if "onboarding was slow" in user:
                return json.dumps({"assignments": [
                    {"code": "Pricing", "quote": "The transparent pricing convinced finance",
                     "confidence": 0.85}]})
            return json.dumps({"assignments": [
                {"code": "Support", "quote": "Support answered within an hour every time",
                 "confidence": 0.95}]})
        if "findings section" in system:
            return json.dumps({"sections": [{"heading": "Overview of Findings", "body": "b"}]})
        return json.dumps({"ok": True})


@pytest.fixture()
def mock_llm(monkeypatch):
    m = MockLLM()
    monkeypatch.setattr(llm, "chat", m)
    client.put('/api/settings/keys', json={"anthropic": "sk-test"})
    return m


def make_ca_project(name, docs=(DOC_A, DOC_B, DOC_C)):
    r = client.post('/api/projects', json={
        "name": name, "method": "content_analysis",
        "config": {"provider": "anthropic", "research_question": "q",
                   "ca_mode": "Inductive — derive the codebook from the data"}})
    assert r.status_code == 200, r.text
    pid = r.json()["id"]
    sids = []
    for i, d in enumerate(docs, 1):
        r = client.post(f'/api/projects/{pid}/sources',
                        files={"file": (f"doc_{i}.txt", io.BytesIO(d.encode()), "text/plain")},
                        data={"grp": ""})
        assert r.status_code == 200, r.text
        sids.append(r.json()["id"])
    return pid, sids


def wait_run(run_id, *want, timeout=20):
    for _ in range(timeout * 20):
        d = client.get(f'/api/runs/{run_id}').json()
        if d["status"] in want:
            return d
        time.sleep(0.05)
    raise AssertionError(f"timeout waiting for {want}; last: {d['status']} {d.get('error')}")


def excerpt_count(run_id):
    return db.get_conn().execute(
        "SELECT COUNT(*) c FROM excerpts WHERE run_id=?", (run_id,)).fetchone()["c"]


# ---------- security ----------

def test_spa_traversal_blocked():
    # percent-encoded traversal must never serve files outside frontend/dist
    r = client.get("/%2e%2e/%2e%2e/backend/app/db.py")
    assert b"SQLite persistence layer" not in r.content
    r = client.get("/..%2f..%2fbackend%2fdata%2ftest.db")
    assert b"SQLite format 3" not in r.content


def test_unknown_api_route_is_404_not_index():
    r = client.get("/api/definitely/not/a/route")
    assert r.status_code == 404
    assert "Unknown API route" in r.text


# ---------- run lifecycle ----------

def test_resume_after_failure_does_not_duplicate(mock_llm):
    pid, _ = make_ca_project("Resume Test")
    # fail doc_2's APPLY call once (codebook derivation also contains doc_2
    # text, so the trigger requires the apply-stage marker too)
    mock_llm.fail_once_when = ("CODEBOOK:", "onboarding was slow")
    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    d = wait_run(run_id, "awaiting_review")
    cp = d["pending_checkpoint"]
    client.post(f'/api/runs/{run_id}/checkpoints/{cp["id"]}/resolve', json={"decisions": []})
    d = wait_run(run_id, "failed")
    n_before = excerpt_count(run_id)
    assert n_before >= 2   # doc_1's assignments landed before the failure
    r = client.post(f'/api/runs/{run_id}/resume')
    assert r.status_code == 200
    d = wait_run(run_id, "completed")
    # doc_1: 2 matched (+1 dropped unknown code), doc_2: 1, doc_3: 1 => exactly 4
    assert excerpt_count(run_id) == 4, "resume must not re-code finished segments"
    rep = client.get(f'/api/runs/{run_id}/report').json()
    assert rep["stats"]["total_assignments"] == 4


def test_cancel_stops_spending_mid_stage(mock_llm):
    pid, _ = make_ca_project("Cancel Test")
    state = {"run_id": None, "apply_calls": 0}

    def on_call(user):
        if "CODEBOOK:" not in user:    # only apply-stage prompts carry the codebook
            return
        state["apply_calls"] += 1
        if state["apply_calls"] == 1 and state["run_id"]:
            pipeline.cancel_run(state["run_id"])   # cancel during first apply call
    mock_llm.on_call = on_call

    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    state["run_id"] = run_id
    d = wait_run(run_id, "awaiting_review")
    cp = d["pending_checkpoint"]
    calls_at_resolve = len(mock_llm.calls)
    client.post(f'/api/runs/{run_id}/checkpoints/{cp["id"]}/resolve', json={"decisions": []})
    d = wait_run(run_id, "cancelled")
    time.sleep(0.3)   # let any (wrong) further calls happen
    # exactly one apply call was in flight when cancel hit; the other two
    # sources must never be sent to the model
    assert len(mock_llm.calls) - calls_at_resolve <= 1, \
        "cancellation must stop the stage before the next LLM call"


def test_double_resolve_is_rejected(mock_llm):
    pid, _ = make_ca_project("DoubleResolve Test")
    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    d = wait_run(run_id, "awaiting_review")
    cp = d["pending_checkpoint"]
    r1 = client.post(f'/api/runs/{run_id}/checkpoints/{cp["id"]}/resolve', json={"decisions": []})
    assert r1.status_code == 200
    r2 = client.post(f'/api/runs/{run_id}/checkpoints/{cp["id"]}/resolve', json={"decisions": []})
    assert r2.status_code == 400
    wait_run(run_id, "completed")


def test_reconcile_on_startup():
    conn = db.get_conn()
    pid, rid, sid = db.new_id(), db.new_id(), db.new_id()
    conn.execute("INSERT INTO projects(id,name,method,config,created_at) VALUES(?,?,?,?,?)",
                 (pid, "Orphan", "thematic", "{}", db.now()))
    conn.execute("INSERT INTO runs(id,project_id,status,created_at,updated_at) "
                 "VALUES(?,?,?,?,?)", (rid, pid, "running", db.now(), db.now()))
    conn.execute("INSERT INTO sources(id,project_id,filename,kind,status,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?)",
                 (sid, pid, "a.mp3", "audio", "transcribing", "{}", db.now()))
    conn.commit()
    pipeline.reconcile_on_startup()
    run = conn.execute("SELECT * FROM runs WHERE id=?", (rid,)).fetchone()
    assert run["status"] == "failed" and "Resume" in run["error"]
    src = conn.execute("SELECT * FROM sources WHERE id=?", (sid,)).fetchone()
    assert src["status"] == "error"


def test_run_blocked_while_transcribing(mock_llm):
    pid, _ = make_ca_project("TranscribeBlock Test", docs=(DOC_A,))
    conn = db.get_conn()
    conn.execute("INSERT INTO sources(id,project_id,filename,kind,status,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?)",
                 (db.new_id(), pid, "b.mp3", "audio", "transcribing", "{}", db.now()))
    conn.commit()
    r = client.post(f'/api/projects/{pid}/runs')
    assert r.status_code == 400 and "transcribing" in r.text


# ---------- checkpoint decision semantics ----------

def _mini_run(method="thematic"):
    """A project+run with codes/excerpts inserted directly, for unit-level
    resolution tests."""
    conn = db.get_conn()
    pid, rid = db.new_id(), db.new_id()
    conn.execute("INSERT INTO projects(id,name,method,config,created_at) VALUES(?,?,?,?,?)",
                 (pid, "Mini", method, json.dumps({"provider": "anthropic"}), db.now()))
    sid = db.new_id()
    conn.execute("INSERT INTO sources(id,project_id,filename,kind,status,text,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?,?)",
                 (sid, pid, "s.txt", "text", "ready", DOC_A, "{}", db.now()))
    conn.execute("INSERT INTO runs(id,project_id,status,created_at,updated_at) "
                 "VALUES(?,?,?,?,?)", (rid, pid, "running", db.now(), db.now()))
    conn.commit()
    project = db.row_to_dict(conn.execute("SELECT * FROM projects WHERE id=?", (pid,)).fetchone(), ("config",))
    sources = [db.row_to_dict(conn.execute("SELECT * FROM sources WHERE id=?", (sid,)).fetchone(), ("meta",))]
    ctx = RunContext(rid, project, sources, project["config"], "anthropic", "m", "k")
    return ctx, sid


def test_merge_chain_and_theme_children_survive():
    ctx, sid = _mini_run()
    a = ctx.add_code("code a", "", "open_code")
    b = ctx.add_code("code b", "", "open_code")
    c = ctx.add_code("code c", "", "open_code")
    ctx.add_excerpt(a, sid, "the price was transparent")
    ctx.add_excerpt(b, sid, "Support was responsive")
    ctx.add_excerpt(c, sid, "the price was transparent")
    t1 = ctx.add_code("Theme 1", "", "theme")
    t2 = ctx.add_code("Theme 2", "", "theme")
    conn = db.get_conn()
    conn.execute("UPDATE codes SET parent_id=? WHERE id IN (?,?)", (t1, a, b))
    conn.execute("UPDATE codes SET parent_id=? WHERE id=?", (t2, c))
    conn.commit()

    # merge chain within one batch: b -> a while a -> c... a merged into c
    apply_code_review_resolution(ctx, {"decisions": [
        {"id": b, "action": "merge", "merge_into": a},
        {"id": a, "action": "merge", "merge_into": c},
    ], "stage": "open_code"})
    rows = {r["id"]: dict(r) for r in conn.execute(
        "SELECT * FROM codes WHERE run_id=?", (ctx.run_id,)).fetchall()}
    assert rows[b]["status"] == "merged" and rows[a]["status"] == "merged"
    # all three excerpts must end on c (the chain terminus), none stranded
    n_on_c = conn.execute("SELECT COUNT(*) n FROM excerpts WHERE code_id=?", (c,)).fetchone()["n"]
    assert n_on_c == 3

    # merging theme t1 into t2 must re-parent its children
    apply_code_review_resolution(ctx, {"decisions": [
        {"id": t1, "action": "merge", "merge_into": t2}], "stage": "theme"})
    kids_of_t2 = [r["id"] for r in conn.execute(
        "SELECT id FROM codes WHERE parent_id=?", (t2,)).fetchall()]
    assert set(kids_of_t2) >= {a, b, c}


def test_deleted_theme_children_land_in_uncategorized():
    ctx, sid = _mini_run()
    k = ctx.add_code("orphan code", "", "open_code")
    ctx.add_excerpt(k, sid, "the price was transparent")
    t = ctx.add_code("Doomed theme", "", "theme")
    conn = db.get_conn()
    conn.execute("UPDATE codes SET parent_id=? WHERE id=?", (t, k))
    conn.commit()
    apply_code_review_resolution(ctx, {"decisions": [{"id": t, "action": "delete"}],
                                       "stage": "theme"})
    common.assemble_report(ctx, "T", [], "theme", "open_code")
    payload = json.loads(conn.execute(
        "SELECT payload FROM reports WHERE run_id=?", (ctx.run_id,)).fetchone()["payload"])
    unc = [th for th in payload["themes"] if th["name"] == "Uncategorized"]
    assert unc and unc[0]["children"][0]["id"] == k, \
        "evidence under a deleted theme must surface in an Uncategorized bucket"


def test_rename_marks_user_edited_and_define_name_respects_it(monkeypatch):
    ctx, sid = _mini_run()
    t = ctx.add_code("Machine name", "", "theme")
    apply_code_review_resolution(ctx, {"decisions": [
        {"id": t, "action": "rename", "name": "Researcher's Name"}], "stage": "theme"})
    code = ctx.codes(stage="theme")[0]
    assert code["meta"].get("user_edited") is True

    from app.methods import thematic
    def fake_chat(*a, **k):
        raise AssertionError("define_name must not call the model when every theme is researcher-named")
    monkeypatch.setattr(llm, "chat", fake_chat)
    thematic.stage_define_name(ctx)   # must no-op, not raise
    assert ctx.codes(stage="theme")[0]["name"] == "Researcher's Name"


# ---------- deletion & cleanup ----------

def test_delete_source_with_excerpts(mock_llm):
    pid, sids = make_ca_project("DeleteSource Test", docs=(DOC_A,))
    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    d = wait_run(run_id, "awaiting_review")
    client.post(f'/api/runs/{run_id}/checkpoints/{d["pending_checkpoint"]["id"]}/resolve',
                json={"decisions": []})
    wait_run(run_id, "completed")
    assert excerpt_count(run_id) > 0
    # a completed run cites this source: refused unless forced, so a report
    # never silently loses the documents it quotes
    r = client.delete(f'/api/sources/{sids[0]}')
    assert r.status_code == 409 and "completed run" in r.json()["detail"]
    r = client.delete(f'/api/sources/{sids[0]}?force=true')
    assert r.status_code == 200, "forced source deletion must not 500 on FK constraints"
    conn = db.get_conn()
    assert conn.execute("SELECT COUNT(*) c FROM sources WHERE id=?", (sids[0],)).fetchone()["c"] == 0
    assert conn.execute("SELECT COUNT(*) c FROM excerpts WHERE source_id=?", (sids[0],)).fetchone()["c"] == 0
    assert not list(db.UPLOADS_DIR.glob(f"{sids[0]}_*")), "uploaded file must be removed"


def test_update_project_config(mock_llm):
    pid, _ = make_ca_project("Update Test", docs=(DOC_A,))
    r = client.put(f'/api/projects/{pid}', json={
        "name": "Updated Name", "method": "thematic",
        "config": {"provider": "anthropic", "research_question": "new q"}})
    assert r.status_code == 200
    p = client.get(f'/api/projects/{pid}').json()
    assert p["name"] == "Updated Name" and p["method"] == "thematic"
    assert p["config"]["research_question"] == "new q"


def test_docx_unicode_project_name():
    conn = db.get_conn()
    rid = db.new_id()
    pid = db.new_id()
    conn.execute("INSERT INTO projects(id,name,method,config,created_at) VALUES(?,?,?,?,?)",
                 (pid, "Étude café", "thematic", "{}", db.now()))
    conn.execute("INSERT INTO runs(id,project_id,status,created_at,updated_at) VALUES(?,?,?,?,?)",
                 (rid, pid, "completed", db.now(), db.now()))
    payload = {"title": "T", "method": "thematic", "project_name": "Étude café",
               "generated_at": db.now(), "provider": "anthropic", "model": "m",
               "sources": [], "sections": [], "themes": [], "stats": {},
               "audit": {"events": 0, "checkpoints": [], "usage": {}}}
    conn.execute("INSERT INTO reports(run_id,payload,created_at) VALUES(?,?,?)",
                 (rid, json.dumps(payload), db.now()))
    conn.commit()
    r = client.get(f'/api/runs/{rid}/report.docx')
    assert r.status_code == 200
    assert "filename*=UTF-8''" in r.headers["content-disposition"]


# ---------- llm layer ----------

def test_truncated_json_is_refused(monkeypatch):
    def truncated_chat(*a, **k):
        return '{"codes": [', {"input_tokens": 1, "output_tokens": 1,
                               "stop_reason": llm.TRUNCATED}
    monkeypatch.setattr(llm, "chat", truncated_chat)
    with pytest.raises(llm.LLMError, match="truncated"):
        llm.chat_json("anthropic", "m", "k", "s", "u")


def test_extract_json_falls_past_bad_fence():
    text = 'Here:\n```json\nnot json at all\n```\nBut also {"ok": true} trailing'
    assert llm._extract_json(text) == {"ok": True}


def test_locate_quote_typographic_normalization():
    text = "She said “I don’t know — really” and left."
    start, end = locate_quote(text, 'I don\'t know - really')
    assert start is not None and text[start:end] == "I don’t know — really"


# ---------- round-2 fixes (from the fix-verification pass) ----------

def test_cyclic_merge_canonicalizes_instead_of_stranding():
    ctx, sid = _mini_run()
    a = ctx.add_code("code a", "", "open_code")
    b = ctx.add_code("code b", "", "open_code")
    ctx.add_excerpt(a, sid, "the price was transparent")
    ctx.add_excerpt(b, sid, "Support was responsive")
    apply_code_review_resolution(ctx, {"decisions": [
        {"id": a, "action": "merge", "merge_into": b},
        {"id": b, "action": "merge", "merge_into": a},
    ], "stage": "open_code"})
    conn = db.get_conn()
    active = [r["id"] for r in conn.execute(
        "SELECT id FROM codes WHERE run_id=? AND status='active'", (ctx.run_id,)).fetchall()]
    assert len(active) == 1, "a merge cycle must canonicalize on exactly one kept code"
    n = conn.execute("SELECT COUNT(*) n FROM excerpts WHERE code_id=?",
                     (active[0],)).fetchone()["n"]
    assert n == 2, "all evidence must land on the kept code"


def test_self_merge_is_kept_not_deleted():
    ctx, sid = _mini_run()
    a = ctx.add_code("code a", "", "open_code")
    ctx.add_excerpt(a, sid, "the price was transparent")
    apply_code_review_resolution(ctx, {"decisions": [
        {"id": a, "action": "merge", "merge_into": a}], "stage": "open_code"})
    row = db.get_conn().execute("SELECT status FROM codes WHERE id=?", (a,)).fetchone()
    assert row["status"] == "active", "self-merge must be a no-op, never a delete"


def test_failed_resolution_reopens_checkpoint(mock_llm, monkeypatch):
    pid, _ = make_ca_project("ReopenCP Test", docs=(DOC_A,))
    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    d = wait_run(run_id, "awaiting_review")
    cp = d["pending_checkpoint"]
    method = pipeline.METHODS["content_analysis"]
    stage = next(s for s in method.stages if s.name == cp["stage"])
    original = stage.apply_resolution

    def boom(ctx, resolution):
        raise RuntimeError("induced apply failure")
    monkeypatch.setattr(stage, "apply_resolution", boom)
    quiet = TestClient(app, base_url="http://127.0.0.1", headers=AUTH,
                       raise_server_exceptions=False)
    r = quiet.post(f'/api/runs/{run_id}/checkpoints/{cp["id"]}/resolve',
                   json={"decisions": []})
    assert r.status_code >= 500
    # the checkpoint must be pending again and the run still reviewable
    d = client.get(f'/api/runs/{run_id}').json()
    assert d["status"] == "awaiting_review"
    assert d["pending_checkpoint"] and d["pending_checkpoint"]["id"] == cp["id"], \
        "a failed resolution must reopen the checkpoint, not wedge the run"
    monkeypatch.setattr(stage, "apply_resolution", original)
    r = client.post(f'/api/runs/{run_id}/checkpoints/{cp["id"]}/resolve',
                    json={"decisions": []})
    assert r.status_code == 200
    wait_run(run_id, "completed")


def test_reconcile_repairs_orphaned_review_state():
    conn = db.get_conn()
    pid, rid = db.new_id(), db.new_id()
    conn.execute("INSERT INTO projects(id,name,method,config,created_at) VALUES(?,?,?,?,?)",
                 (pid, "OrphanCP", "thematic", "{}", db.now()))
    conn.execute("INSERT INTO runs(id,project_id,status,created_at,updated_at) "
                 "VALUES(?,?,?,?,?)", (rid, pid, "awaiting_review", db.now(), db.now()))
    conn.commit()   # awaiting_review but NO pending checkpoint
    pipeline.reconcile_on_startup()
    run = conn.execute("SELECT * FROM runs WHERE id=?", (rid,)).fetchone()
    assert run["status"] == "failed" and "Resume" in run["error"]


def test_familiarize_rejects_malformed_output_without_poisoning_state(monkeypatch):
    ctx, sid = _mini_run()
    def list_chat(provider, model, api_key, system, user, **kw):
        return json.dumps([{"summary": "s", "memo": "m"}, "junk"]), \
               {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"}
    monkeypatch.setattr(llm, "chat", list_chat)
    # a salvageable list-wrapped object is unwrapped, not rejected
    common.stage_familiarize(ctx)
    assert isinstance(ctx.state["summaries"][sid], dict)

    ctx2, sid2 = _mini_run()
    def junk_chat(provider, model, api_key, system, user, **kw):
        return json.dumps(["just", "strings"]), \
               {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"}
    monkeypatch.setattr(llm, "chat", junk_chat)
    with pytest.raises(RuntimeError, match="Familiarization"):
        common.stage_familiarize(ctx2)
    assert sid2 not in ctx2.state.get("summaries", {}), \
        "malformed output must not be persisted into resumable state"


def test_put_project_method_change_blocked_after_runs(mock_llm):
    pid, _ = make_ca_project("PutGuard Test", docs=(DOC_A,))
    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    d = wait_run(run_id, "awaiting_review")
    # any edit while a run is active is refused
    r = client.put(f'/api/projects/{pid}', json={
        "name": "X", "method": "content_analysis",
        "config": {"provider": "anthropic", "research_question": "q",
                   "ca_mode": "Inductive — derive the codebook from the data"}})
    assert r.status_code == 409
    client.post(f'/api/runs/{run_id}/checkpoints/{d["pending_checkpoint"]["id"]}/resolve',
                json={"decisions": []})
    wait_run(run_id, "completed")
    # method change is refused even after the run finishes
    r = client.put(f'/api/projects/{pid}', json={
        "name": "X", "method": "thematic",
        "config": {"provider": "anthropic", "research_question": "q"}})
    assert r.status_code == 409


def test_delete_source_blocked_while_run_active(mock_llm):
    pid, sids = make_ca_project("DeleteBusy Test", docs=(DOC_A,))
    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    d = wait_run(run_id, "awaiting_review")
    r = client.delete(f'/api/sources/{sids[0]}')
    assert r.status_code == 409
    client.post(f'/api/runs/{run_id}/checkpoints/{d["pending_checkpoint"]["id"]}/resolve',
                json={"decisions": []})
    wait_run(run_id, "completed")


def test_locate_quote_double_quote_class_not_dead():
    text = 'He called it “a total mess” at the time.'
    start, end = locate_quote(text, 'called it "a total mess" at')
    assert start is not None and text[start:end] == 'called it “a total mess” at'


def test_oversized_malformed_json_is_refused_not_sliced(monkeypatch):
    big = "not json " * 10000   # > 60000 chars, unparseable
    def big_chat(provider, model, api_key, system, user, **kw):
        return big, {"input_tokens": 1, "output_tokens": 1, "stop_reason": "end_turn"}
    monkeypatch.setattr(llm, "chat", big_chat)
    with pytest.raises(llm.LLMError, match="too large to repair"):
        llm.chat_json("anthropic", "m", "k", "s", "u")


# ---------- evidence-exploration endpoints (checkpoint panel + reader) ----------

def test_code_excerpts_endpoint(mock_llm):
    pid, sids = make_ca_project("Evidence Test", docs=(DOC_A,))
    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    d = wait_run(run_id, "awaiting_review")
    client.post(f'/api/runs/{run_id}/checkpoints/{d["pending_checkpoint"]["id"]}/resolve',
                json={"decisions": []})
    wait_run(run_id, "completed")
    conn = db.get_conn()
    code = conn.execute("SELECT id FROM codes WHERE run_id=? AND name='Pricing'",
                        (run_id,)).fetchone()
    r = client.get(f'/api/runs/{run_id}/codes/{code["id"]}/excerpts')
    assert r.status_code == 200
    exs = r.json()
    assert len(exs) == 1 and exs[0]["source"] == "doc_1.txt"
    assert exs[0]["quote"] == "the price was transparent"
    assert exs[0]["start_char"] is not None


def test_coded_source_endpoint(mock_llm):
    pid, sids = make_ca_project("Reader Test", docs=(DOC_A,))
    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    d = wait_run(run_id, "awaiting_review")
    client.post(f'/api/runs/{run_id}/checkpoints/{d["pending_checkpoint"]["id"]}/resolve',
                json={"decisions": []})
    wait_run(run_id, "completed")
    r = client.get(f'/api/runs/{run_id}/sources/{sids[0]}/coded')
    assert r.status_code == 200
    data = r.json()
    assert data["filename"] == "doc_1.txt" and data["text"] == DOC_A
    assert len(data["spans"]) == 2          # Pricing + Support (unknown code dropped)
    for s in data["spans"]:
        assert data["text"][s["start"]:s["end"]] == s["quote"], \
            "span offsets must reproduce the exact quote"
    names = {c["name"] for c in data["codes"]}
    assert names == {"Pricing", "Support"}
    # a deleted code's evidence must not appear in the reader
    conn = db.get_conn()
    code = conn.execute("SELECT id FROM codes WHERE run_id=? AND name='Pricing'",
                        (run_id,)).fetchone()
    conn.execute("UPDATE codes SET status='deleted' WHERE id=?", (code["id"],))
    conn.commit()
    data = client.get(f'/api/runs/{run_id}/sources/{sids[0]}/coded').json()
    assert {c["name"] for c in data["codes"]} == {"Support"}


def test_coded_source_wrong_project_404(mock_llm):
    pid1, sids1 = make_ca_project("ReaderA", docs=(DOC_A,))
    pid2, _ = make_ca_project("ReaderB", docs=(DOC_B,))
    run2 = client.post(f'/api/projects/{pid2}/runs').json()["run_id"]
    d = wait_run(run2, "awaiting_review")
    # source from project 1 must not be readable through project 2's run
    r = client.get(f'/api/runs/{run2}/sources/{sids1[0]}/coded')
    assert r.status_code == 404
    client.post(f'/api/runs/{run2}/checkpoints/{d["pending_checkpoint"]["id"]}/resolve',
                json={"decisions": []})
    wait_run(run2, "completed")


# ---------- round-3 fixes (reader/checkpoint verification pass) ----------

def test_coded_source_utf16_offsets_with_emoji():
    """Python indexes by code point, JS by UTF-16 unit — offsets shipped to
    the reader must be JS-compatible or every highlight after an emoji shifts."""
    ctx, sid = _mini_run()
    conn = db.get_conn()
    text = "Intro 🎉🎉 said: the app is great. More 🚀 text follows here."
    conn.execute("UPDATE sources SET text=? WHERE id=?", (text, sid))
    conn.commit()
    ctx.sources[0]["text"] = text
    c = ctx.add_code("praising the app", "", "open_code")
    ctx.add_excerpt(c, sid, "the app is great")
    ctx.add_excerpt(c, sid, "text follows")
    data = client.get(f'/api/runs/{ctx.run_id}/sources/{sid}/coded').json()
    assert len(data["spans"]) == 2
    utf16 = data["text"].encode("utf-16-le")   # emulate a JS string
    for s in data["spans"]:
        sliced = utf16[2 * s["start"]:2 * s["end"]].decode("utf-16-le")
        assert sliced == s["quote"], f"JS slice {sliced!r} != quote {s['quote']!r}"


def test_code_excerpts_includes_grouping_children():
    """A theme's evidence lives on its child codes; the evidence panel must
    show it rather than 'No excerpts'."""
    ctx, sid = _mini_run()
    child = ctx.add_code("child code", "", "open_code")
    ctx.add_excerpt(child, sid, "the price was transparent")
    theme = ctx.add_code("Big Theme", "", "theme")
    conn = db.get_conn()
    conn.execute("UPDATE codes SET parent_id=? WHERE id=?", (theme, child))
    conn.commit()
    exs = client.get(f'/api/runs/{ctx.run_id}/codes/{theme}/excerpts').json()
    assert len(exs) == 1 and exs[0]["via"] == "child code"
    # a direct excerpt has no 'via'
    exs = client.get(f'/api/runs/{ctx.run_id}/codes/{child}/excerpts').json()
    assert len(exs) == 1 and exs[0]["via"] is None


def test_reader_legend_counts_split_located_unlocated():
    ctx, sid = _mini_run()
    c = ctx.add_code("mixed code", "", "open_code")
    ctx.add_excerpt(c, sid, "the price was transparent")       # locatable
    ctx.add_excerpt(c, sid, "utterly absent paraphrase")       # not in text
    data = client.get(f'/api/runs/{ctx.run_id}/sources/{sid}/coded').json()
    meta = data["codes"][0]
    assert meta["count"] == 1 and meta["unlocated_count"] == 1
    assert len(data["spans"]) == 1 and len(data["unlocated"]) == 1


def test_cleared_definition_is_honored_blank_name_is_not():
    ctx, _ = _mini_run()
    c = ctx.add_code("a code", "old definition", "open_code")
    apply_code_review_resolution(ctx, {"decisions": [
        {"id": c, "action": "rename", "name": "", "definition": ""}], "stage": "open_code"})
    row = db.get_conn().execute("SELECT name, definition FROM codes WHERE id=?", (c,)).fetchone()
    assert row["name"] == "a code", "blank name must keep the old name"
    assert row["definition"] == "", "blank definition is a deliberate clearing"


# ---------- visualizations ----------

def test_gt_report_carries_model_stats_and_docx_embeds_figure(mock_llm, monkeypatch):
    # reuse the CA mock answers where possible; supply GT-specific ones
    base = mock_llm._answer
    def gt_answer(system, user):
        if "reading data closely" in system:
            return json.dumps({"summary": "s", "memo": "m", "notable_features": []})
        if "OPEN CODING" in system:
            return json.dumps({"codes": [
                {"name": "valuing transparency", "definition": "d",
                 "excerpts": [{"quote": "the price was transparent", "memo": ""}]},
                {"name": "trusting support", "definition": "d",
                 "excerpts": [{"quote": "Support was responsive", "memo": ""}]}]})
        if "AXIAL" in system or "grouping open codes" in system:
            import re as _re
            ids = _re.findall(r"\[([0-9a-f]{12})\]", user)
            return json.dumps({"categories": [
                {"name": "Building vendor trust", "definition": "d", "rationale": "r",
                 "code_ids": ids}]})
        if "SELECTIVE CODING" in system:
            return json.dumps({"core_category": {"name": "Earning trust", "definition": "d",
                                                 "is_existing_category_id": None},
                               "storyline": "story", "relationships": [],
                               "theoretical_gaps": []})
        return base(system, user)
    monkeypatch.setattr(mock_llm, "_answer", gt_answer)

    r = client.post('/api/projects', json={
        "name": "GT Viz Test", "method": "grounded_theory",
        "config": {"provider": "anthropic", "research_question": "q",
                   "gt_variant": "Straussian (axial coding with paradigm model)"}})
    pid = r.json()["id"]
    client.post(f'/api/projects/{pid}/sources',
                files={"file": ("doc_1.txt", io.BytesIO(DOC_A.encode()), "text/plain")},
                data={"grp": ""})
    run_id = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    for _ in range(3):   # three checkpoints
        d = wait_run(run_id, "awaiting_review", "completed")
        if d["status"] == "completed":
            break
        client.post(f'/api/runs/{run_id}/checkpoints/{d["pending_checkpoint"]["id"]}/resolve',
                    json={"decisions": []})
    wait_run(run_id, "completed")
    rep = client.get(f'/api/runs/{run_id}/report').json()
    assert rep["stats"]["kind"] == "gt_model"
    assert rep["stats"]["core"]["name"] == "Earning trust"
    # docx must embed the model figure as an image part
    docx_bytes = client.get(f'/api/runs/{run_id}/report.docx').content
    import zipfile, io as _io
    names = zipfile.ZipFile(_io.BytesIO(docx_bytes)).namelist()
    assert any(n.startswith("word/media/") for n in names), \
        "grounded theory docx must embed the model figure"


def test_viz_renderers_handle_empty_data():
    from app import viz
    assert viz.gt_model_png({"themes": [], "stats": {}}) is None
    assert viz.thematic_map_png({"themes": []}) is None
    assert viz.content_freq_png({"rows": []}) is None
    assert viz.framework_heatmap_png({"rows": [], "codes": []}) is None
    assert viz.render_for_payload({"method": "thematic", "themes": [], "stats": {}}) == []


def test_gt_relation_labels_only_core_targets():
    """A relationship aimed at another category must NOT be labeled on that
    category's arrow into the core — that would assert a relationship the
    analysis never made."""
    from app.viz import gt_relation_labels
    stats = {"core": {"name": "Core", "existing_category_id": "cat9"},
             "relationships": [
                 {"from_category_id": "cat1", "relation": "condition for", "to": "core"},
                 {"from_category_id": "cat2", "relation": "condition for", "to": "cat5"},
                 {"from_category_id": "cat3", "relation": "strategy for", "to": "cat9"},
                 {"from_category_id": "cat1", "relation": "context of", "to": "CORE"},
                 {"from_category_id": "cat4", "relation": "consequence of", "to": ""}]}
    labels = gt_relation_labels(stats)
    assert labels["cat1"] == "condition for / context of"
    assert "cat2" not in labels, "category-to-category relation must not be labeled"
    assert labels["cat3"] == "strategy for", "core's own category id counts as core"
    assert labels["cat4"] == "consequence of", "missing target defaults to core"


def test_gt_model_excludes_core_category_satellite_and_caps():
    from app import viz
    themes = [{"id": f"c{i}", "name": f"Category {i}",
               "children": [], "excerpts": [{}] * i} for i in range(14)]
    payload = {"themes": themes,
               "stats": {"kind": "gt_model",
                         "core": {"name": "Core", "existing_category_id": "c3"},
                         "relationships": []}}
    png = viz.gt_model_png(payload)
    assert png and png[:8] == b"\x89PNG\r\n\x1a\n"
    # if ALL categories are the core/Uncategorized there is nothing to draw
    assert viz.gt_model_png({"themes": [{"id": "x", "name": "Uncategorized"}],
                             "stats": {"core": {"existing_category_id": None}}}) is None


def test_package_script_manifest_and_relative_path():
    import subprocess, tempfile, zipfile, pathlib as pl
    root = pl.Path(__file__).resolve().parent.parent.parent
    with tempfile.TemporaryDirectory() as td:
        r = subprocess.run(["bash", str(root / "package.sh"), "bundle.zip"],
                           cwd=td, capture_output=True, text=True)
        assert r.returncode == 0, r.stderr
        out = pl.Path(td) / "bundle.zip"
        assert out.exists(), "relative output path must land in the caller's cwd"
        names = zipfile.ZipFile(out).namelist()
        joined = "\n".join(names)
        for forbidden in ("qualilens.db", ".venv", "node_modules", "Publication",
                          "Sample Transcripts", ".pytest_cache",
                          ".DS_Store", "backend/data", "FINGERPRINT", "devnotes"):
            assert forbidden not in joined, f"bundle must not contain {forbidden}"
        for required in ("QualiLens/run.sh", "QualiLens/LICENSE", "QualiLens/NOTICE",
                         "QualiLens/backend/app/main.py",
                         "QualiLens/backend/app/methods/literature_synthesis.py",
                         "QualiLens/frontend/dist/index.html",
                         "QualiLens/frontend/dist/manual.html",
                         "QualiLens/frontend/public/manual.html",
                         "QualiLens/manual/build_manual.py",
                         "QualiLens/manual/01-getting-started.md",
                         "QualiLens/manual/08-literature-synthesis.md",
                         "QualiLens/backend/app/models.json",
                         "QualiLens/VERSION", "QualiLens/RELEASE"):
            assert required in names, f"bundle must contain {required}"
        # RELEASE names the changelog's top entry, or says the bundle is interim
        import re as _re
        top = _re.search(r"^## (\S+)", (root / "CHANGELOG.md").read_text(), _re.M)
        expect = top.group(1) if top and _re.fullmatch(r"\d+\.\d+\.\d+", top.group(1)) else "unreleased"
        assert zipfile.ZipFile(out).read("QualiLens/RELEASE").decode().strip() == expect


# ---------- provenance marks ----------

def test_provenance_marks_present():
    conn = db.get_conn()
    app_id = conn.execute("PRAGMA application_id").fetchone()[0]
    assert app_id == db.APPLICATION_ID, \
        "every database the app creates must carry the QLns application_id"
    assert db.get_setting("lineage") == db.LINEAGE
    # the docx generator identifies itself in the document properties
    import docx as _docx
    from io import BytesIO as _B
    from app import report_docx
    payload = {"title": "T", "method": "thematic", "project_name": "X",
               "generated_at": db.now(), "provider": "p", "model": "m",
               "sources": [], "sections": [], "themes": [], "stats": {},
               "audit": {"events": 0, "checkpoints": [], "usage": {}}}
    d = _docx.Document(_B(report_docx.build_docx(payload)))
    assert "QualiLens" in (d.core_properties.comments or "")


# ---------- model catalog maintenance ----------

def test_model_catalog_loads_from_json_with_fallback(tmp_path, monkeypatch):
    # the shipped catalog must parse and drive PROVIDERS
    assert llm.CATALOG_PATH.exists()
    assert llm.PROVIDERS["anthropic"]["models"], "catalog must list models"
    assert llm.DEFAULT_PRICES["anthropic"][1] > llm.DEFAULT_PRICES["anthropic"][0]
    # a broken file must fall back, never raise
    bad = tmp_path / "models.json"
    bad.write_text("{not json")
    monkeypatch.setattr(llm, "CATALOG_PATH", bad)
    cat = llm._load_catalog()
    assert set(cat) == set(llm._FALLBACK_CATALOG)
    assert cat["openai"]["default_model"] in cat["openai"]["models"]
    # a default not present in models must be repaired to the first model
    bad.write_text(json.dumps({"anthropic": {"models": ["m1", "m2"],
                                             "default_model": "retired-model"}}))
    cat = llm._load_catalog()
    assert cat["anthropic"]["default_model"] == "m1"


def test_check_models_endpoint_diffs_catalog(monkeypatch):
    client.put('/api/settings/keys', json={"anthropic": "sk-test"})
    retired = llm.PROVIDERS["anthropic"]["models"][0]
    live = [m for m in llm.PROVIDERS["anthropic"]["models"][1:]] + ["claude-new-99"]

    def fake_list(provider, key):
        assert provider == "anthropic" and key
        return live
    monkeypatch.setattr(llm, "list_models", fake_list)
    r = client.post('/api/settings/check_models', json={"provider": "anthropic"})
    assert r.status_code == 200
    res = r.json()["anthropic"]
    assert res["ok"] and retired in res["missing"]
    entry = next(c for c in res["catalog"] if c["id"] == retired)
    assert entry["available"] is False
    assert "claude-new-99" in res["live"]
    # a provider without a key reports, not errors
    client.put('/api/settings/keys', json={"mistral": "__clear__"})
    r = client.post('/api/settings/check_models', json={"provider": "mistral"})
    assert r.json()["mistral"]["ok"] is False


def test_retired_model_error_gets_diagnosis(monkeypatch):
    def gone(*a, **k):
        raise llm.LLMError("HTTP 404: {'type': 'error', 'error': {'type': "
                           "'not_found_error', 'message': 'model: nope-1'}}")
    monkeypatch.setattr(llm, "_chat_impl", gone)
    with pytest.raises(llm.LLMError, match="retired or renamed"):
        llm.chat("anthropic", "nope-1", "sk-x", "s", "u")
    # unrelated errors pass through untouched
    def boom(*a, **k):
        raise llm.LLMError("HTTP 500: overloaded")
    monkeypatch.setattr(llm, "_chat_impl", boom)
    with pytest.raises(llm.LLMError, match="overloaded") as ei:
        llm.chat("anthropic", "nope-1", "sk-x", "s", "u")
    assert "retired" not in str(ei.value)


# ---------- catalog verification round ----------

def test_load_catalog_survives_wrong_shape(tmp_path, monkeypatch):
    bad = tmp_path / "models.json"
    monkeypatch.setattr(llm, "CATALOG_PATH", bad)
    # valid JSON, wrong shapes — none may raise, all must fall back sanely
    for content in ('["a", "list"]', '"just a string"', '17',
                    json.dumps({"anthropic": {"models": "not-a-list"}})):
        bad.write_text(content)
        cat = llm._load_catalog()
        assert set(cat) == set(llm._FALLBACK_CATALOG)
        assert cat["anthropic"]["models"] == llm._FALLBACK_CATALOG["anthropic"]["models"]


def test_check_models_reads_catalog_fresh_from_disk(tmp_path, monkeypatch):
    """A maintainer edits models.json and presses Check — the diff must use
    the file as it is NOW, not the import-time snapshot."""
    edited = tmp_path / "models.json"
    edited.write_text(json.dumps({"anthropic": {
        "models": ["brand-new-model"], "default_model": "brand-new-model"}}))
    monkeypatch.setattr(llm, "CATALOG_PATH", edited)
    monkeypatch.setattr(llm, "list_models", lambda p, k: ["brand-new-model"])
    client.put('/api/settings/keys', json={"anthropic": "sk-test"})
    r = client.post('/api/settings/check_models', json={"provider": "anthropic"})
    res = r.json()["anthropic"]
    assert [c["id"] for c in res["catalog"]] == ["brand-new-model"]
    assert res["missing"] == []


def test_check_models_rejects_unknown_provider():
    r = client.post('/api/settings/check_models', json={"provider": "nonsense"})
    assert r.status_code == 400


def test_project_model_is_trimmed(mock_llm):
    r = client.post('/api/projects', json={
        "name": "Trim Test", "method": "thematic",
        "config": {"provider": "anthropic", "research_question": "q",
                   "model": "  gpt-x-typed-with-spaces \n"}})
    assert r.status_code == 200
    assert r.json()["config"]["model"] == "gpt-x-typed-with-spaces"


def test_mistral_invalid_model_gets_diagnosis(monkeypatch):
    def mistral_400(*a, **k):
        raise llm.LLMError('HTTP 400: {"object": "error", "message": '
                           '"Invalid model: mistral-medium-2312", '
                           '"type": "invalid_model", "code": "1500"}')
    monkeypatch.setattr(llm, "_chat_impl", mistral_400)
    with pytest.raises(llm.LLMError, match="retired or renamed"):
        llm.chat("mistral", "mistral-medium-2312", "sk-x", "s", "u")


# ---------- in-place app update ----------

# a throwaway release key for the tests; the real one never ships
from app import signing as _signing
_TEST_SEED = bytes(range(32))
TEST_PUBLIC_HEX = _signing.public_hex_from_seed(_TEST_SEED)


def _make_bundle(tmp_path, extra=None, notice=True, version="9999.01.01",
                 sign=True, seed=_TEST_SEED, tamper_after_sign=None):
    """A minimal bundle. Signed with the test key by default (the updater is
    monkeypatched to trust it); sign=False yields an unsigned bundle,
    tamper_after_sign={member: content} alters files after signing."""
    import hashlib as _hl
    import zipfile as _zf
    zpath = tmp_path / "bundle.zip"
    files = {
        "QualiLens/run.sh": "#!/bin/bash\necho new-version\n",
        "QualiLens/backend/app/main.py": "# new main\n",
        "QualiLens/frontend/dist/index.html": "<!doctype html>new",
        "QualiLens/VERSION": version,
    }
    if notice:
        files["QualiLens/NOTICE"] = "QualiLens\nCopyright 2026 Ashita Aggarwal and Suraj Commuri\n"
    files.update(extra or {})
    if sign:
        lines = []
        for name in sorted(files):
            data = files[name].encode() if isinstance(files[name], str) else files[name]
            lines.append(f"{_hl.sha256(data).hexdigest()}  {name[len('QualiLens/'):]}")
        manifest = "\n".join(lines) + "\n"
        files["QualiLens/MANIFEST.sha256"] = manifest
        files["QualiLens/MANIFEST.sig"] = _signing.sign_bytes(manifest.encode(), seed) + "\n"
    files.update(tamper_after_sign or {})
    with _zf.ZipFile(zpath, "w") as z:
        for name, content in files.items():
            z.writestr(name, content)
    return zpath


def _fake_app_root(tmp_path, monkeypatch):
    from app import update
    root = tmp_path / "installed"
    (root / "backend" / "app").mkdir(parents=True)
    (root / "backend" / "data").mkdir(parents=True)
    (root / "frontend" / "dist").mkdir(parents=True)
    (root / "backend" / "app" / "main.py").write_text("# old main")
    (root / "backend" / "app" / "old_module.py").write_text("# should disappear")
    (root / "backend" / "data" / "qualilens.db").write_text("PRECIOUS USER DATA")
    (root / "frontend" / "dist" / "stale-asset.js").write_text("stale")
    (root / "run.sh").write_text("#!/bin/bash\necho old")
    (root / "VERSION").write_text("1111.01.01")
    monkeypatch.setattr(update, "APP_ROOT", root)
    monkeypatch.setattr(update, "BACKUP_DIR", root / ".update-backup")
    monkeypatch.setattr(update, "PUBLIC_KEY_HEX", TEST_PUBLIC_HEX)
    return root


def test_update_replaces_app_and_never_touches_data(tmp_path, monkeypatch):
    from app import update
    root = _fake_app_root(tmp_path, monkeypatch)
    r = update.apply_update(_make_bundle(tmp_path))
    assert r["ok"] and r["from_version"] == "1111.01.01" and r["to_version"] == "9999.01.01"
    assert (root / "backend" / "app" / "main.py").read_text() == "# new main\n"
    # user data is byte-identical
    assert (root / "backend" / "data" / "qualilens.db").read_text() == "PRECIOUS USER DATA"
    # replaced trees do not keep files deleted upstream
    assert not (root / "backend" / "app" / "old_module.py").exists()
    assert not (root / "frontend" / "dist" / "stale-asset.js").exists()
    # run.sh stays executable; a backup of the old version exists
    assert (root / "run.sh").stat().st_mode & 0o111
    assert any((root / ".update-backup").rglob("main.py"))


def test_update_refuses_zip_slip_and_data_paths(tmp_path, monkeypatch):
    from app import update
    root = _fake_app_root(tmp_path, monkeypatch)
    # a bundle trying to write outside the app root is refused outright
    evil = _make_bundle(tmp_path, extra={"QualiLens/../../etc/evil": "x"})
    with pytest.raises(update.UpdateError, match="Unsafe path"):
        update.apply_update(evil)
    # a bundle carrying a data file installs the app but REFUSES the data path
    sneaky = _make_bundle(tmp_path, extra={
        "QualiLens/backend/data/qualilens.db": "ATTACKER DATA",
        "QualiLens/devnotes/notes.md": "should not install"})
    r = update.apply_update(sneaky)
    assert (root / "backend" / "data" / "qualilens.db").read_text() == "PRECIOUS USER DATA"
    assert not (root / "devnotes").exists()
    assert "backend/data/qualilens.db" in r["files_refused"]


def test_update_refuses_foreign_bundles(tmp_path, monkeypatch):
    from app import update
    _fake_app_root(tmp_path, monkeypatch)
    with pytest.raises(update.UpdateError, match="not look like a QualiLens bundle"):
        update.apply_update(_make_bundle(tmp_path, notice=False))
    not_zip = tmp_path / "x.zip"
    not_zip.write_text("hello")
    with pytest.raises(update.UpdateError, match="not a zip"):
        update.apply_update(not_zip)


def _quiesce_runs():
    """Updates are refused while any run is live; earlier tests leave scratch
    runs in 'running'. Park them so update tests exercise the updater."""
    conn = db.get_conn()
    conn.execute("UPDATE runs SET status='cancelled' WHERE status IN ('running','awaiting_review')")
    conn.commit()


def test_update_endpoint_applies_and_reports(tmp_path, monkeypatch):
    from app import update
    root = _fake_app_root(tmp_path, monkeypatch)
    _quiesce_runs()
    bundle = _make_bundle(tmp_path)
    with open(bundle, "rb") as f:
        r = client.post('/api/settings/update',
                        files={"file": ("QualiLens.zip", f, "application/zip")})
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["ok"] and body["to_version"] == "9999.01.01"
    assert "restart_required" not in body  # QUALILENS_TEST suppresses self-stop
    assert (root / "backend" / "data" / "qualilens.db").read_text() == "PRECIOUS USER DATA"


# ---------- literature synthesis ----------

def test_join_pdf_pages_records_original_page_numbers():
    from app import ingestion
    text, pages = ingestion.join_pdf_pages(["", "First real page", "", "  ", "Last page"])
    assert text == "First real page\n\nLast page"
    assert pages == [{"page": 2, "start": 0, "end": 15},
                     {"page": 5, "start": 17, "end": 26}]
    for p in pages:
        assert text[p["start"]:p["end"]].strip()
    assert ingestion.page_for_offset(pages, 0) == 2
    assert ingestion.page_for_offset(pages, 17) == 5
    assert ingestion.page_for_offset(pages, 16) is None   # inside the join
    assert ingestion.page_for_offset(pages, 999) is None
    assert ingestion.page_for_offset([], 3) is None
    assert ingestion.page_for_offset(pages, None) is None
    # a PDF with no extractable text yields no page map at all
    t2, p2 = ingestion.join_pdf_pages(["", "  "])
    assert t2 == "" and p2 is None


def test_citation_guard_flags_only_out_of_corpus():
    from app.methods.literature_synthesis import _citation_guard
    labels = ["Okafor, 2021", "van der Berg, 2019"]
    body = ("Trust recurs (Okafor, 2021). Sampling is thin (van der Berg, 2019). "
            "Okafor (2021) also notes fees. A classic view persists "
            "(Smith, 1998; Jones, 2003). As was shown long before now (2010).")
    flagged = _citation_guard([{"heading": "S", "body": body}], labels)
    assert "Smith, 1998" in flagged and "Jones, 2003" in flagged
    assert any("2010" in f for f in flagged)
    assert all("Okafor" not in f and "Berg" not in f for f in flagged)
    # a bare year right after a corpus surname is that paper, not a phantom
    assert _citation_guard([{"body": "Okafor (2021) argues this."}], labels) == []
    # a phantom packed beside a real citation must still be caught — each
    # ';'-separated citation is judged on its own
    assert _citation_guard([{"body": "Both agree (Okafor, 2021; Smith, 1998)."}],
                           labels) == ["Smith, 1998"]
    # two-letter surnames count as corpus vocabulary
    assert _citation_guard([{"body": "Trust matters (Li, 2020)."}],
                           ["Li, 2020"]) == []
    # the same phantom twice flags once
    assert _citation_guard([{"body": "(Smith, 1998) and again (Smith, 1998)."}],
                           labels) == ["Smith, 1998"]
    assert _citation_guard([], labels) == []


def test_concept_matrix_figure_and_empty_safe():
    from app import viz
    stats = {"kind": "concept_matrix",
             "codes": ["Transparency", "Trust"],
             "rows": [{"source": "Alpha, 2021",
                       "cells": {"Transparency": {"n": 2, "summary": "s"},
                                 "Trust": {"n": 0, "summary": ""}}},
                      {"source": "Beta, 2022",
                       "cells": {"Transparency": {"n": 1, "summary": "s"},
                                 "Trust": {"n": 3, "summary": "s"}}}]}
    figs = viz.render_for_payload({"method": "literature_synthesis", "stats": stats})
    assert len(figs) == 1
    caption, png = figs[0]
    assert "Concept-by-paper" in caption and png[:8] == b"\x89PNG\r\n\x1a\n"
    assert viz.render_for_payload(
        {"method": "literature_synthesis",
         "stats": {"kind": "concept_matrix", "rows": [], "codes": []}}) == []


def test_extraction_review_apply_marks_user_edited_and_reapplies_as_noop():
    from app.methods import literature_synthesis as ls
    ctx, sid = _mini_run("literature_synthesis")
    ctx.state["extractions"] = {sid: {"label": "Alpha, 2021", "citation": "c",
                                      "aims": "a", "method": "m", "sample": "s",
                                      "findings": "f", "limitations": "l",
                                      "user_edited": []}}
    # excluding EVERY paper must refuse before mutating anything — the
    # checkpoint reopens instead of stranding the run at synthesis
    with pytest.raises(ValueError, match="re-include at least one"):
        ls.cp_extraction_apply(ctx, {"rows": [{"source_id": sid, "exclude": True,
                                               "findings": "should not stick"}]})
    assert ctx.state["extractions"][sid]["findings"] == "f"

    res = {"rows": [{"source_id": sid, "findings": " Edited findings. ",
                     "label": "", "exclude": False}]}
    ls.cp_extraction_apply(ctx, res)
    row = ctx.state["extractions"][sid]
    assert row["findings"] == "Edited findings."
    assert row["label"] == "Alpha, 2021"          # a blank label is invalid; kept
    assert "findings" in row["user_edited"] and "label" not in row["user_edited"]
    assert ctx.state["excluded"][sid] is False
    # apply runs on a fresh ctx at resolution time, so it must persist itself
    saved = json.loads(db.get_conn().execute(
        "SELECT state FROM runs WHERE id=?", (ctx.run_id,)).fetchone()["state"])
    assert saved["extractions"][sid]["findings"] == "Edited findings."
    ls.cp_extraction_apply(ctx, res)              # a re-apply is a no-op
    assert ctx.state["extractions"][sid]["findings"] == "Edited findings."
    ls.cp_extraction_apply(ctx, {"rows": [{"source_id": "nope", "aims": "x"}]})


def test_synthesis_refuses_ungrounded_concepts(monkeypatch):
    """The never-cites-from-memory principle, enforced structurally: support
    referencing unknown excerpt ids is dropped, and a synthesis whose every
    concept is ungrounded fails the stage instead of fabricating a review."""
    from app.methods import literature_synthesis as ls
    ctx, sid = _mini_run("literature_synthesis")
    ctx.config["research_question"] = "q"
    codes = ls._field_codes(ctx)
    eid = ctx.add_excerpt(codes["findings"], sid, "the price was transparent")
    ctx.state["extractions"] = {sid: {"label": "Alpha, 2021", "citation": "",
                                      "aims": "", "method": "", "sample": "",
                                      "findings": "F", "limitations": "",
                                      "user_edited": []}}

    answers = {"support": [{"excerpt_id": "ffffffffffff", "point": "from memory"}]}

    def fake(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        return (json.dumps({"concepts": [
            {"name": "Phantom", "definition": "d", "rationale": "r",
             "support": answers["support"]}]}),
            {"input_tokens": 1, "output_tokens": 1})

    monkeypatch.setattr(llm, "chat", fake)
    with pytest.raises(RuntimeError, match="corpus-grounded"):
        ls.stage_synthesize(ctx)
    assert ctx.codes(stage="concept") == []

    answers["support"] = [{"excerpt_id": eid, "point": "grounded"},
                          {"excerpt_id": "ffffffffffff", "point": "from memory"}]
    ls.stage_synthesize(ctx)
    concepts = ctx.codes(stage="concept")
    assert len(concepts) == 1 and concepts[0]["name"] == "Phantom"
    exs = ctx.excerpts_for(concepts[0]["id"])
    assert len(exs) == 1 and exs[0]["quote"] == "the price was transparent"
    assert exs[0]["start_char"] is not None       # provenance re-locates the quote


def test_synthesis_bars_unlocated_quotes_and_refuses_empty_pool(monkeypatch):
    """An extraction quote that cannot be located verbatim in its paper must
    not ground a concept; and a corpus with no located quotes at all must
    refuse BEFORE spending the synthesis call."""
    from app.methods import literature_synthesis as ls
    ctx, sid = _mini_run("literature_synthesis")
    ctx.config["research_question"] = "q"
    codes = ls._field_codes(ctx)
    # this quote does not occur in DOC_A -> stored unlocated (start_char NULL)
    fake_eid = ctx.add_excerpt(codes["findings"], sid, "a fabricated sentence xyzzy")
    row = db.get_conn().execute("SELECT start_char FROM excerpts WHERE id=?",
                                (fake_eid,)).fetchone()
    assert row["start_char"] is None
    ctx.state["extractions"] = {sid: {"label": "Alpha, 2021", "citation": "",
                                      "aims": "", "method": "", "sample": "",
                                      "findings": "F", "limitations": "",
                                      "user_edited": []}}
    calls = []

    def fake(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        calls.append(user)
        return json.dumps({"concepts": []}), {"input_tokens": 1, "output_tokens": 1}

    monkeypatch.setattr(llm, "chat", fake)
    with pytest.raises(RuntimeError, match="No located extraction quotes"):
        ls.stage_synthesize(ctx)
    assert calls == [], "the synthesis call must not be spent on an empty pool"

    # with one located quote, the unlocated one is still not offered
    good_eid = ctx.add_excerpt(codes["findings"], sid, "the price was transparent")

    def fake2(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        assert fake_eid not in user and good_eid in user
        return (json.dumps({"concepts": [
            {"name": "C", "definition": "d", "rationale": "r",
             "support": [{"excerpt_id": good_eid, "point": "ok"},
                         {"excerpt_id": fake_eid, "point": "unlocated"}]}]}),
            {"input_tokens": 1, "output_tokens": 1})

    monkeypatch.setattr(llm, "chat", fake2)
    ls.stage_synthesize(ctx)
    concepts = ctx.codes(stage="concept")
    exs = ctx.excerpts_for(concepts[0]["id"])
    assert len(exs) == 1 and exs[0]["quote"] == "the price was transparent"


def test_synthesis_merges_duplicate_concept_names(monkeypatch):
    from app.methods import literature_synthesis as ls
    ctx, sid = _mini_run("literature_synthesis")
    ctx.config["research_question"] = "q"
    codes = ls._field_codes(ctx)
    e1 = ctx.add_excerpt(codes["findings"], sid, "the price was transparent")
    e2 = ctx.add_excerpt(codes["findings"], sid, "Support was responsive")
    ctx.state["extractions"] = {sid: {"label": "Alpha, 2021", "citation": "",
                                      "aims": "", "method": "", "sample": "",
                                      "findings": "F", "limitations": "",
                                      "user_edited": []}}

    def fake(provider, model, api_key, system, user, max_tokens=8000, temperature=0.3):
        return (json.dumps({"concepts": [
            {"name": "Trust", "definition": "d1", "rationale": "r",
             "support": [{"excerpt_id": e1, "point": "a"}]},
            {"name": "trust", "definition": "d2", "rationale": "r",
             "support": [{"excerpt_id": e2, "point": "b"}]}]}),
            {"input_tokens": 1, "output_tokens": 1})

    monkeypatch.setattr(llm, "chat", fake)
    ls.stage_synthesize(ctx)
    concepts = ctx.codes(stage="concept")
    assert len(concepts) == 1, "same-named concepts must merge, not duplicate"
    assert len(ctx.excerpts_for(concepts[0]["id"])) == 2


def test_extraction_apply_refuses_label_collision():
    from app.methods import literature_synthesis as ls
    ctx, sid = _mini_run("literature_synthesis")
    sid2 = db.new_id()
    conn = db.get_conn()
    conn.execute("INSERT INTO sources(id,project_id,filename,kind,status,text,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?,?)",
                 (sid2, ctx.project["id"], "t.txt", "text", "ready", DOC_B, "{}", db.now()))
    conn.commit()
    ctx.sources.append(db.row_to_dict(
        conn.execute("SELECT * FROM sources WHERE id=?", (sid2,)).fetchone(), ("meta",)))
    ctx.state["extractions"] = {
        sid: {"label": "Alpha, 2021", "citation": "", "aims": "", "method": "",
              "sample": "", "findings": "f", "limitations": "", "user_edited": []},
        sid2: {"label": "Beta, 2022", "citation": "", "aims": "", "method": "",
               "sample": "", "findings": "f", "limitations": "", "user_edited": []}}
    with pytest.raises(ValueError, match="share the label"):
        ls.cp_extraction_apply(ctx, {"rows": [{"source_id": sid2,
                                               "label": "alpha, 2021"}]})
    assert ctx.state["extractions"][sid2]["label"] == "Beta, 2022"  # untouched


def test_upload_blocked_mid_run_and_corpus_frozen(mock_llm):
    """A source uploaded while a run exists must not slip into that run's
    analysis: the upload API refuses during running/awaiting_review, and the
    run's corpus is frozen at start for the resume path."""
    pid, sids = make_ca_project("FrozenCorpus", docs=(DOC_A,))
    r = client.post(f'/api/projects/{pid}/runs')
    assert r.status_code == 200
    rid = r.json()["run_id"]
    d = wait_run(rid, "awaiting_review")
    # 1) the endpoint refuses mid-run uploads outright
    r = client.post(f'/api/projects/{pid}/sources',
                    files={"file": ("late.txt", io.BytesIO(b"late data"), "text/plain")},
                    data={"grp": ""})
    assert r.status_code == 409 and "run is in progress" in r.json()["detail"]
    # 2) even a source that appears in the DB anyway (older app, direct write)
    #    stays out of the frozen corpus on resume/continuation
    late = db.new_id()
    conn = db.get_conn()
    conn.execute("INSERT INTO sources(id,project_id,filename,kind,status,text,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?,?)",
                 (late, pid, "late.txt", "text", "ready", "late data", "{}", db.now()))
    conn.commit()
    cp = d["pending_checkpoint"]
    r = client.post(f'/api/runs/{rid}/checkpoints/{cp["id"]}/resolve',
                    json={"decisions": []})
    assert r.status_code == 200, r.text
    wait_run(rid, "completed")
    rep = client.get(f'/api/runs/{rid}/report').json()
    assert [s["filename"] for s in rep["sources"]] == ["doc_1.txt"], \
        "a late source must not join a frozen run"


def test_docx_renders_concept_matrix_extraction_table_and_page_anchors():
    import docx as _docx
    from io import BytesIO as _B
    from app import report_docx
    payload = {"title": "T", "method": "literature_synthesis", "project_name": "X",
               "generated_at": db.now(), "provider": "p", "model": "m",
               "sources": [{"id": "s1", "filename": "a.pdf", "grp": None}],
               "sections": [{"heading": "The Corpus", "body": "Alpha, 2021 (a.pdf)"}],
               "themes": [{"id": "c1", "name": "Transparency", "definition": "d",
                           "children": [],
                           "excerpts": [{"quote": "q", "memo": "", "source": "a.pdf",
                                         "source_id": "s1", "page": 3}]}],
               "stats": {"kind": "concept_matrix", "codes": ["Transparency"],
                         "rows": [{"source": "Alpha, 2021",
                                   "cells": {"Transparency": {"n": 1, "summary": "sum"}}}],
                         "extraction_rows": [{"source_id": "s1", "filename": "a.pdf",
                                              "label": "Alpha, 2021", "citation": "cite",
                                              "fields": {"aims": "A"}, "excluded": False}],
                         "field_labels": {"aims": "Aims"}},
               "source_summaries": [],
               "audit": {"events": 0, "checkpoints": [], "usage": {}}}
    d = _docx.Document(_B(report_docx.build_docx(payload)))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "Concept-by-Paper Matrix" in text
    assert "Transparency (1): sum" in text
    assert "Appendix: Extraction Table" in text and "Aims: A" in text
    assert ", p. 3" in text                       # page anchor beside the quote
    assert "Evidence: Concepts and Excerpts" in text
    assert "Themes, Codes, and Evidence" not in text


# ---------- pull-only release updates ----------

def test_check_updates_compares_build_stamps(monkeypatch):
    from app import update
    rel = {"tag_name": "v9.9.9", "name": "v9.9.9 — build 9999.01.01-0000",
           "body": "", "html_url": "https://github.com/iconerr/qualilens/releases/tag/v9.9.9",
           "assets": [{"name": "QualiLens.zip", "size": 5_000_000,
                       "browser_download_url":
                           "https://github.com/iconerr/qualilens/releases/download/v9.9.9/QualiLens.zip"}]}
    monkeypatch.setattr(update, "fetch_latest_release", lambda: rel)
    body = client.post('/api/settings/check_updates').json()
    assert body["ok"] and body["newer"] and body["has_bundle"]
    assert body["build"] == "9999.01.01-0000" and body["tag"] == "v9.9.9"
    # the same build as installed is not an update
    rel2 = dict(rel, name=f"v0.0.1 — build {update._current_version()}")
    monkeypatch.setattr(update, "fetch_latest_release", lambda: rel2)
    assert client.post('/api/settings/check_updates').json()["newer"] is False
    # a release with no build stamp is reported honestly, never as newer
    rel3 = dict(rel, name="v1.2.3", body="notes")
    monkeypatch.setattr(update, "fetch_latest_release", lambda: rel3)
    b3 = client.post('/api/settings/check_updates').json()
    assert b3["newer"] is False and "build stamp" in b3["note"]
    # a network failure is a clean answer, not a 500
    def boom():
        raise update.UpdateError("Could not reach GitHub: no route")
    monkeypatch.setattr(update, "fetch_latest_release", boom)
    b4 = client.post('/api/settings/check_updates').json()
    assert b4["ok"] is False and "GitHub" in b4["error"]


def test_install_update_pulls_release_and_applies(tmp_path, monkeypatch):
    """The release path reuses apply_update wholesale, so user data survives
    by the same allowlist; the release is resolved server-side only."""
    from app import update
    root = _fake_app_root(tmp_path, monkeypatch)
    _quiesce_runs()
    bundle = _make_bundle(tmp_path)
    monkeypatch.setattr(update, "download_latest_bundle", lambda dest_dir: bundle)
    r = client.post('/api/settings/install_update')
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["ok"] and body["to_version"] == "9999.01.01"
    assert "restart_required" not in body   # QUALILENS_TEST suppresses self-stop
    assert (root / "backend" / "data" / "qualilens.db").read_text() == "PRECIOUS USER DATA"
    def refuse(dest_dir):
        raise update.UpdateError("The latest release carries no QualiLens.zip asset.")
    monkeypatch.setattr(update, "download_latest_bundle", refuse)
    r = client.post('/api/settings/install_update')
    assert r.status_code == 400 and "asset" in r.json()["detail"]


def test_download_refuses_foreign_asset_host(monkeypatch):
    import pathlib as pl
    from app import update
    rel = {"assets": [{"name": "QualiLens.zip", "size": 10,
                       "browser_download_url": "https://evil.example.com/QualiLens.zip"}]}
    monkeypatch.setattr(update, "fetch_latest_release", lambda: rel)
    with pytest.raises(update.UpdateError, match="not hosted"):
        update.download_latest_bundle(pl.Path("/tmp"))


# ---------- branching: revisit a review as a new run ----------

def test_branch_run_reopens_checkpoint_and_recomputes(mock_llm):
    pid, sids = make_ca_project("BranchCA")
    rid = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    d = wait_run(rid, "awaiting_review")
    client.post(f"/api/runs/{rid}/checkpoints/{d['pending_checkpoint']['id']}/resolve",
                json={"decisions": []})
    wait_run(rid, "completed")
    conn = db.get_conn()
    src_codes = conn.execute("SELECT COUNT(*) c FROM codes WHERE run_id=?",
                             (rid,)).fetchone()["c"]
    src_ex = excerpt_count(rid)

    r = client.post(f'/api/runs/{rid}/branch', json={"stage": "review_codebook"})
    assert r.status_code == 200, r.text
    nid = r.json()["run_id"]
    d2 = wait_run(nid, "awaiting_review")
    assert d2["pending_checkpoint"]["stage"] == "review_codebook"
    # codes are copied under FRESH ids, fully disjoint from the source run's
    new_ids = {x["id"] for x in conn.execute(
        "SELECT id FROM codes WHERE run_id=?", (nid,)).fetchall()}
    old_ids = {x["id"] for x in conn.execute(
        "SELECT id FROM codes WHERE run_id=?", (rid,)).fetchall()}
    assert new_ids and not (new_ids & old_ids)
    # the coding stage runs AFTER this review, so its output is left behind:
    # the branch reopens the codebook with no excerpts, and re-codes after
    assert excerpt_count(nid) == 0
    client.post(f"/api/runs/{nid}/checkpoints/{d2['pending_checkpoint']['id']}/resolve",
                json={"decisions": []})
    wait_run(nid, "completed")
    assert client.get(f'/api/runs/{nid}/report').json()["stats"]["kind"] \
        == "content_frequencies"
    assert excerpt_count(nid) == src_ex, "the branch re-coded the corpus afresh"
    # the source run is untouched — same codes, same evidence, report intact
    assert conn.execute("SELECT COUNT(*) c FROM codes WHERE run_id=?",
                        (rid,)).fetchone()["c"] == src_codes
    assert excerpt_count(rid) == src_ex
    assert client.get(f'/api/runs/{rid}/report').status_code == 200


def test_branch_refusals(mock_llm):
    pid, _ = make_ca_project("BranchRefuse", docs=(DOC_A,))
    rid = client.post(f'/api/projects/{pid}/runs').json()["run_id"]
    d = wait_run(rid, "awaiting_review")
    r = client.post(f'/api/runs/{rid}/branch', json={"stage": "review_codebook"})
    assert r.status_code == 400 and "waiting at this review" in r.json()["detail"]
    r = client.post(f'/api/runs/{rid}/branch', json={"stage": "apply"})
    assert r.status_code == 400 and "review checkpoints" in r.json()["detail"]
    r = client.post(f'/api/runs/{rid}/branch', json={"stage": "nope"})
    assert r.status_code == 400
    client.post(f"/api/runs/{rid}/checkpoints/{d['pending_checkpoint']['id']}/resolve",
                json={"decisions": []})
    wait_run(rid, "completed")


def test_branch_prunes_later_stage_state(mock_llm):
    """A branch must not inherit artifacts of stages after the branch point —
    a stale matrix would let the matrix stage skip recomputation and ignore
    the researcher's revised decisions."""
    conn = db.get_conn()
    pid, rid = db.new_id(), db.new_id()
    conn.execute("INSERT INTO projects(id,name,method,config,created_at) VALUES(?,?,?,?,?)",
                 (pid, "FW", "framework",
                  json.dumps({"provider": "anthropic", "research_question": "q",
                              "codebook_text": "A: a"}), db.now()))
    sid = db.new_id()
    conn.execute("INSERT INTO sources(id,project_id,filename,kind,status,text,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?,?)",
                 (sid, pid, "s.txt", "text", "ready", DOC_A, "{}", db.now()))
    state = {"source_ids": [sid], "done_units": [f"chart:{sid}:0"],
             "matrix_rows": {sid: {"A": {"summary": "stale", "n": 1}}}}
    conn.execute("INSERT INTO runs(id,project_id,status,stage_index,state,created_at,updated_at) "
                 "VALUES(?,?,?,?,?,?,?)",
                 (rid, pid, "completed", 4, json.dumps(state), db.now(), db.now()))
    conn.commit()
    nid = pipeline.branch_run(rid, "review_charting")
    wait_run(nid, "awaiting_review")
    st = json.loads(conn.execute("SELECT state FROM runs WHERE id=?",
                                 (nid,)).fetchone()["state"])
    assert "matrix_rows" not in st, "later-stage artifacts must not survive a branch"
    assert st["done_units"] == state["done_units"], "work before the review is kept"
    assert st["source_ids"] == [sid], "the frozen corpus is inherited"


def test_gt_bucket_places_paradigm_roles():
    from app.viz import gt_bucket
    assert gt_bucket("situational conditions for") == "conditions"
    assert gt_bucket("sensory basis for") == "conditions"
    assert gt_bucket("contextual amplifier of") == "context"
    assert gt_bucket("cognitive framing that supports") == "context"
    assert gt_bucket("core dimension of") == "dimensions"
    assert gt_bucket("self-regulation strategies that sustain") == "strategies"
    assert gt_bucket("consequence of") == "consequences"
    assert gt_bucket("") == "related"
    assert gt_bucket("relates to") == "related"


def test_branch_remaps_code_ids_inside_meta(mock_llm):
    """A GT branch at review_core must remap the category ids the core code's
    meta carries (relationships, is_existing_category_id) — stale ids would
    silently corrupt the branched report's paradigm figure."""
    conn = db.get_conn()
    pid, rid = db.new_id(), db.new_id()
    conn.execute("INSERT INTO projects(id,name,method,config,created_at) VALUES(?,?,?,?,?)",
                 (pid, "GT", "grounded_theory",
                  json.dumps({"provider": "anthropic", "research_question": "q"}),
                  db.now()))
    sid = db.new_id()
    conn.execute("INSERT INTO sources(id,project_id,filename,kind,status,text,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?,?)",
                 (sid, pid, "s.txt", "text", "ready", DOC_A, "{}", db.now()))
    cat, core = db.new_id(), db.new_id()
    state = {"source_ids": [sid], "core_id": core}
    conn.execute("INSERT INTO runs(id,project_id,status,stage_index,state,created_at,updated_at) "
                 "VALUES(?,?,?,?,?,?,?)",
                 (rid, pid, "completed", 8, json.dumps(state), db.now(), db.now()))
    conn.execute("INSERT INTO codes(id,run_id,name,definition,stage,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?)",
                 (cat, rid, "Cat", "", "category", "{}", db.now()))
    core_meta = {"storyline": "s", "is_existing_category_id": cat,
                 "relationships": [{"from_category_id": cat,
                                    "relation": "condition for", "to": "core"}]}
    conn.execute("INSERT INTO codes(id,run_id,name,definition,stage,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?)",
                 (core, rid, "Core", "", "core", json.dumps(core_meta), db.now()))
    conn.commit()
    nid = pipeline.branch_run(rid, "review_core")
    wait_run(nid, "awaiting_review")
    new_cat = conn.execute("SELECT id FROM codes WHERE run_id=? AND stage='category'",
                           (nid,)).fetchone()["id"]
    new_core = db.row_to_dict(conn.execute(
        "SELECT * FROM codes WHERE run_id=? AND stage='core'", (nid,)).fetchone(),
        ("meta",))
    assert new_cat != cat
    assert new_core["meta"]["is_existing_category_id"] == new_cat
    assert new_core["meta"]["relationships"][0]["from_category_id"] == new_cat
    assert new_core["meta"]["relationships"][0]["to"] == "core"   # literal kept
    st = json.loads(conn.execute("SELECT state FROM runs WHERE id=?",
                                 (nid,)).fetchone()["state"])
    assert st["core_id"] == new_core["id"]
    assert st["branched_from"] == rid and st["branched_at"] == "review_core"
    cp = conn.execute("SELECT instructions FROM checkpoints WHERE run_id=? "
                      "AND status='pending'", (nid,)).fetchone()
    assert "already applied" in cp["instructions"], \
        "the reopened review must say its earlier decisions are baked in"


def test_branch_backfills_corpus_freeze_for_legacy_runs(mock_llm):
    """A source run from before the corpus freeze has no source_ids snapshot;
    its branch must not inherit that openness — the corpus is reconstructed
    from what the copied artifacts reference."""
    conn = db.get_conn()
    pid, rid = db.new_id(), db.new_id()
    conn.execute("INSERT INTO projects(id,name,method,config,created_at) VALUES(?,?,?,?,?)",
                 (pid, "Legacy", "content_analysis",
                  json.dumps({"provider": "anthropic", "research_question": "q",
                              "ca_mode": "Deductive — I will supply the codebook",
                              "codebook_text": "A: a"}), db.now()))
    sid = db.new_id()
    conn.execute("INSERT INTO sources(id,project_id,filename,kind,status,text,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?,?)",
                 (sid, pid, "orig.txt", "text", "ready", DOC_A, "{}", db.now()))
    conn.execute("INSERT INTO runs(id,project_id,status,stage_index,state,created_at,updated_at) "
                 "VALUES(?,?,?,?,?,?,?)",
                 (rid, pid, "completed", 4, "{}", db.now(), db.now()))
    cid = db.new_id()
    conn.execute("INSERT INTO codes(id,run_id,name,definition,stage,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?)", (cid, rid, "A", "a", "codebook", "{}", db.now()))
    conn.execute("INSERT INTO excerpts(id,run_id,code_id,source_id,quote,memo,created_at) "
                 "VALUES(?,?,?,?,?,?,?)",
                 (db.new_id(), rid, cid, sid, "the price was transparent", "", db.now()))
    # a source added AFTER the legacy run finished must not join its branch
    late = db.new_id()
    conn.execute("INSERT INTO sources(id,project_id,filename,kind,status,text,meta,created_at) "
                 "VALUES(?,?,?,?,?,?,?,?)",
                 (late, pid, "late.txt", "text", "ready", DOC_B, "{}", db.now()))
    conn.commit()
    nid = pipeline.branch_run(rid, "review_codebook")
    wait_run(nid, "awaiting_review")
    st = json.loads(conn.execute("SELECT state FROM runs WHERE id=?",
                                 (nid,)).fetchone()["state"])
    assert st["source_ids"] == [sid], \
        "the branch's corpus is what the copied artifacts reference"
