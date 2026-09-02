# Copyright 2026 Ashita Aggarwal and Suraj Commuri
# SPDX-License-Identifier: Apache-2.0

"""Render a run's report payload to a formatted .docx."""

import time
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x66, 0x66, 0x66)
AMBER_TEXT = RGBColor(0x92, 0x60, 0x0A)


def build_docx(payload: dict) -> bytes:
    doc = Document()
    # generator identification in the document properties (the visible
    # counterpart is the colophon at the end of the report)
    doc.core_properties.comments = "Generated with QualiLens (ql-a2f4467b)"
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    # Title block
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(payload.get("title", "Qualitative Analysis Report"))
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    when = time.strftime("%B %d, %Y", time.localtime(payload.get("generated_at", time.time())))
    r = meta.add_run(
        f"Method: {payload.get('method_label') or payload.get('method', '')}  ·  "
        f"Generated {when} with QualiLens  ·  "
        f"Model: {payload.get('provider', '')}/{payload.get('model', '')}")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    # Method configuration: every setup answer as recorded, verbatim — the
    # option text carries the methodological commitment
    config = payload.get("config") or {}
    if config:
        doc.add_heading("Method Configuration", level=1)
        labels = payload.get("config_labels") or {}
        for key, value in config.items():
            text = str(value if value is not None else "").strip()
            p = doc.add_paragraph()
            rr = p.add_run(f"{labels.get(key, key)}: ")
            rr.bold = True
            p.add_run(text if text else "(blank)")
        p = doc.add_paragraph()
        rr = p.add_run(f"Provider and model: {payload.get('provider', '')}/{payload.get('model', '')}")
        rr.bold = True

    # Sources
    doc.add_heading("Data Sources", level=1)
    for s in payload.get("sources", []):
        line = s["filename"] + (f"  (group: {s['grp']})" if s.get("grp") else "")
        doc.add_paragraph(line, style="List Bullet")

    # Narrative sections
    for sec in payload.get("sections", []):
        doc.add_heading(sec.get("heading", ""), level=1)
        for para in str(sec.get("body", "")).split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # Method-appropriate figure(s): grounded theory model, thematic map,
    # frequency chart, or framework heatmap
    try:
        from . import viz
        figures = viz.render_for_payload(payload)
    except Exception:  # noqa: BLE001 — figures are additive, never blocking
        import logging
        logging.getLogger("qualilens.viz").exception(
            "Figure generation failed; exporting report without figures")
        figures = []
    for caption, png in figures:
        doc.add_picture(BytesIO(png), width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED

    # Frequency table (content analysis)
    stats = payload.get("stats") or {}
    if stats.get("kind") == "content_frequencies" and stats.get("rows"):
        doc.add_heading("Code Frequency Table", level=1)
        if stats.get("unit"):
            p = doc.add_paragraph()
            rr = p.add_run(f"Unit counted: {stats['unit']}. Rates are per 10,000 characters "
                           "of source text and correct for unequal source and group sizes.")
            rr.font.size = Pt(9)
            rr.font.color.rgb = MUTED
        groups = stats.get("groups", [])
        has_rate = any("per_10k_chars" in r for r in stats["rows"])
        cols = 4 + (1 if has_rate else 0) + len(groups) * (2 if has_rate else 1)
        table = doc.add_table(rows=1, cols=cols)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        col = 0
        hdr[col].text = "Code"; col += 1
        hdr[col].text = "Count"; col += 1
        hdr[col].text = "%"; col += 1
        hdr[col].text = "Sources"; col += 1
        if has_rate:
            hdr[col].text = "per 10k chars"; col += 1
        for g in groups:
            hdr[col].text = g; col += 1
            if has_rate:
                hdr[col].text = f"{g} per 10k"; col += 1
        for row in stats["rows"]:
            cells = table.add_row().cells
            col = 0
            cells[col].text = row["code"]; col += 1
            cells[col].text = str(row["count"]); col += 1
            cells[col].text = f"{row['pct']}%"; col += 1
            cells[col].text = str(row.get("sources", "")); col += 1
            if has_rate:
                cells[col].text = str(row.get("per_10k_chars", "")); col += 1
            for g in groups:
                cells[col].text = str(row.get("by_group", {}).get(g, 0)); col += 1
                if has_rate:
                    cells[col].text = str(row.get("by_group_per_10k", {}).get(g, "")); col += 1

    # Framework matrix
    if stats.get("kind") == "framework_matrix" and stats.get("rows"):
        doc.add_heading("Framework Matrix", level=1)
        for row in stats["rows"]:
            doc.add_heading(row["source"], level=2)
            for cname, cell in row["cells"].items():
                if cell.get("n"):
                    p = doc.add_paragraph()
                    r = p.add_run(f"{cname} ({cell['n']}): ")
                    r.bold = True
                    p.add_run(cell.get("summary", ""))

    # Concept-by-paper matrix (literature synthesis) — rendered paper by
    # paper like the framework matrix; a grid wide enough for a full corpus
    # does not survive a page break
    if stats.get("kind") == "concept_matrix" and stats.get("rows"):
        doc.add_heading("Concept-by-Paper Matrix", level=1)
        for row in stats["rows"]:
            doc.add_heading(row["source"], level=2)
            for cname, cell in row["cells"].items():
                if cell.get("n"):
                    p = doc.add_paragraph()
                    r = p.add_run(f"{cname} ({cell['n']}): ")
                    r.bold = True
                    p.add_run(cell.get("summary", ""))

    # Themes with evidence (concepts, for a literature synthesis)
    doc.add_heading("Evidence: Concepts and Excerpts"
                    if payload.get("method") == "literature_synthesis"
                    else "Themes, Codes, and Evidence", level=1)
    for theme in payload.get("themes", []):
        doc.add_heading(theme["name"], level=2)
        if theme.get("definition"):
            p = doc.add_paragraph()
            r = p.add_run(theme["definition"])
            r.italic = True
        _write_excerpts(doc, theme.get("excerpts", []))
        for child in theme.get("children", []):
            doc.add_heading(child["name"], level=3)
            if child.get("definition"):
                p = doc.add_paragraph()
                r = p.add_run(child["definition"])
                r.italic = True
            _write_excerpts(doc, child.get("excerpts", []))

    # Extraction table appendix (literature synthesis)
    if stats.get("kind") == "concept_matrix" and stats.get("extraction_rows"):
        doc.add_heading("Appendix: Extraction Table", level=1)
        field_labels = stats.get("field_labels") or {}
        for row in stats["extraction_rows"]:
            head = row.get("label") or row.get("filename", "")
            if row.get("excluded"):
                head += "  (excluded from the synthesis)"
            doc.add_heading(head, level=2)
            if row.get("citation"):
                p = doc.add_paragraph()
                r = p.add_run(row["citation"])
                r.italic = True
                r.font.size = Pt(10)
            for key, text in (row.get("fields") or {}).items():
                if text:
                    p = doc.add_paragraph()
                    r = p.add_run(f"{field_labels.get(key, key)}: ")
                    r.bold = True
                    p.add_run(text)
            if row.get("cited_work"):
                p = doc.add_paragraph()
                r = p.add_run("Findings the paper attributes to other work (not used in the synthesis): ")
                r.bold = True
                r.font.size = Pt(9)
                r2 = p.add_run(row["cited_work"])
                r2.font.size = Pt(9)
                r2.font.color.rgb = MUTED

    # Familiarization appendix
    summaries = [s for s in payload.get("source_summaries", []) if s.get("summary")]
    if summaries:
        doc.add_heading("Appendix: Source Summaries (Familiarization)", level=1)
        for s in summaries:
            doc.add_heading(s.get("source", ""), level=2)
            doc.add_paragraph(s["summary"])
            if s.get("memo"):
                p = doc.add_paragraph()
                r = p.add_run(f"Analytic memo: {s['memo']}")
                r.italic = True
                r.font.size = Pt(10)

    # Audit appendix
    doc.add_heading("Appendix: Audit Trail", level=1)
    audit = payload.get("audit", {})
    doc.add_paragraph(
        f"This analysis logged {audit.get('events', 0)} events. The complete log — every "
        "model call, every researcher decision with its parameters, and every checkpoint "
        "payload — is exportable from the run screen as a JSON file. Researcher checkpoints:")
    for cp in audit.get("checkpoints", []):
        when_cp = ""
        if cp.get("resolved_at"):
            when_cp = " — resolved " + time.strftime(
                "%Y-%m-%d %H:%M", time.localtime(cp["resolved_at"]))
        doc.add_paragraph(f"{cp.get('title', cp.get('stage'))} ({cp.get('status')}){when_cp}",
                          style="List Bullet")
        summ = cp.get("summary") or {}
        if summ:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            rr = p.add_run(_summary_line(summ))
            rr.font.size = Pt(9)
            rr.font.color.rgb = MUTED
    if audit.get("branched_from"):
        doc.add_paragraph(
            f"This run was branched from run {audit['branched_from']} at "
            f"'{audit.get('branched_at')}'; the work before that review was carried over.")
    usage = audit.get("usage", {})
    if usage:
        doc.add_paragraph(
            f"Model usage: {usage.get('calls', 0)} calls, "
            f"{usage.get('input_tokens', 0):,} input tokens, "
            f"{usage.get('output_tokens', 0):,} output tokens.")
    models_used = audit.get("models_used") or {}
    if models_used:
        doc.add_paragraph("Models that answered, by call: " + ", ".join(
            f"{k} ({v})" for k, v in sorted(models_used.items(), key=lambda kv: -kv[1])) + ".")
    if "excerpts_located" in audit:
        doc.add_paragraph(
            f"Evidence: {audit.get('excerpts_located', 0):,} excerpts located verbatim in their "
            f"sources; {audit.get('excerpts_unlocated', 0):,} could not be located and are "
            "listed as unverified rather than quoted.")

    # Colophon: provenance and a suggested software citation — deliberately
    # NOT a copyright notice, since the report's content belongs entirely to
    # the researchers who produced it (see the project NOTICE file).
    colophon = doc.add_paragraph()
    r = colophon.add_run(
        "This report was generated with QualiLens, free software by Ashita Aggarwal "
        "and Suraj Commuri (Apache-2.0). Suggested software citation: Aggarwal, A., & "
        "Commuri, S. (2026). QualiLens: A local application for LLM-assisted "
        "qualitative data analysis [Computer software]. The content of this report "
        "belongs to its authors.")
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _is_located(e: dict) -> bool:
    if "located" in e:
        return bool(e["located"])
    if "start_char" not in e and "end_char" not in e:
        return True        # a payload from before the flag existed: no verdict, quote it
    return e.get("start_char") is not None and e.get("end_char") is not None


def _summary_line(summ: dict) -> str:
    bits = []
    acts = summ.get("decisions") or {}
    if acts:
        bits.append("decisions: " + ", ".join(f"{k} {v}" for k, v in sorted(acts.items())))
    if summ.get("renamed_to"):
        bits.append("renamed to: " + "; ".join(str(x) for x in summ["renamed_to"][:12])
                    + ("…" if len(summ["renamed_to"]) > 12 else ""))
    if summ.get("added"):
        bits.append("added: " + "; ".join(str(x) for x in summ["added"][:12]))
    if summ.get("excerpts_removed"):
        bits.append(f"excerpts removed: {summ['excerpts_removed']}")
    if summ.get("extraction_rows_edited"):
        bits.append(f"extraction rows edited: {summ['extraction_rows_edited']}")
    if summ.get("papers_excluded"):
        bits.append(f"papers excluded: {summ['papers_excluded']}")
    return " · ".join(bits) if bits else "approved without changes"


def _write_excerpts(doc, excerpts: list, limit: int = 12) -> None:
    """Located excerpts are quoted; excerpts whose quote could not be found
    verbatim in the source are listed afterwards as UNVERIFIED, in a
    different register, never inside quotation marks."""
    located = [e for e in excerpts if _is_located(e)]
    unlocated = [e for e in excerpts if not _is_located(e)]
    for e in located[:limit]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        r = p.add_run(f"“{e['quote']}”")
        r.font.size = Pt(10)
        where = f"  — {e.get('source', '')}"
        if e.get("page"):
            where += f", p. {e['page']}"   # PDF page anchor, when known
        src = p.add_run(where)
        src.font.size = Pt(9)
        src.font.color.rgb = MUTED
        if e.get("memo"):
            m = doc.add_paragraph()
            m.paragraph_format.left_indent = Inches(0.6)
            mr = m.add_run(f"Memo: {e['memo']}")
            mr.font.size = Pt(9)
            mr.font.color.rgb = MUTED
    if len(located) > limit:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        r = p.add_run(f"(+{len(located) - limit} further located excerpts in the project database)")
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
    if unlocated:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        r = p.add_run(f"Unverified ({len(unlocated)}): the model returned text under this code "
                      "that could not be found verbatim in the source — likely paraphrase. "
                      "Not quoted; check against the document before use.")
        r.font.size = Pt(9)
        r.font.color.rgb = AMBER_TEXT
        for e in unlocated[:limit]:
            q = doc.add_paragraph()
            q.paragraph_format.left_indent = Inches(0.6)
            t = str(e.get("quote", ""))
            rr = q.add_run("[not located verbatim] " + (t[:300] + ("…" if len(t) > 300 else "")))
            rr.font.size = Pt(9)
            rr.italic = True
            rr.font.color.rgb = MUTED
            w = q.add_run(f"  — {e.get('source', '')}")
            w.font.size = Pt(9)
            w.font.color.rgb = MUTED
        if len(unlocated) > limit:
            q = doc.add_paragraph()
            q.paragraph_format.left_indent = Inches(0.6)
            rr = q.add_run(f"(+{len(unlocated) - limit} further unverified excerpts in the project database)")
            rr.font.size = Pt(9)
            rr.font.color.rgb = MUTED
